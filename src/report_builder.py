"""Report builders for the AI-FDS report generator page."""

from __future__ import annotations

import html
import io
import json
from collections import Counter
from typing import Literal

from src.time_utils import formatted_now, now_for_app


DEFAULT_RECOMMENDATION = (
    "Verify suspicious contact through official channels, do not share OTP/passwords, "
    "preserve the evidence, and report urgent financial or identity-related requests "
    "to the relevant campus or organisation support team."
)

DEFAULT_SECTIONS = {
    "summary": True,
    "evidence": True,
    "explanations": True,
    "risk": True,
    "recommendations": True,
    "appendix": True,
}


def _text(value: object, fallback: str = "-") -> str:
    if value is None:
        return fallback
    value_text = str(value).strip()
    return value_text or fallback


def _percent(value: object) -> float:
    try:
        return max(0.0, min(100.0, float(value)))
    except (TypeError, ValueError):
        return 0.0


def _date_text(value: object) -> str:
    text = _text(value)
    return text.replace("T", " ")[:19]


def _flags(value: object) -> list[str]:
    if value is None:
        return []
    if isinstance(value, str):
        try:
            return _flags(json.loads(value))
        except json.JSONDecodeError:
            return [value] if value.strip() else []
    if isinstance(value, list):
        return [_text(item, "") for item in value if _text(item, "")]
    return [_text(value)]


def _risk_bucket(row: dict[str, object]) -> str:
    prediction = _text(row.get("prediction"), "").lower()
    confidence = _percent(row.get("confidence"))
    if any(term in prediction for term in ("legitimate", "lower risk", "real human", "benign", "safe")):
        return "Lower risk"
    if any(term in prediction for term in ("suspicious", "scam", "phishing", "ai-generated", "high risk", "chunk")):
        return "High risk" if confidence >= 65 else "Needs review"
    if confidence >= 75:
        return "High risk"
    if confidence >= 40:
        return "Needs review"
    return "Lower risk"


def _risk_counts(rows: list[dict[str, object]]) -> Counter:
    return Counter(_risk_bucket(row) for row in rows)


def _prediction_color(prediction: object) -> str:
    text = _text(prediction, "").lower()
    if any(term in text for term in ("lower risk", "legitimate", "real human", "safe")):
        return "#0891B2"
    if any(term in text for term in ("suspicious", "scam", "phishing", "ai-generated", "high risk", "chunk")):
        return "#DC2626"
    return "#D97706"


def _confidence_chart_png(rows: list[dict[str, object]]) -> bytes | None:
    """Render the report confidence overview as PNG for PDF/DOCX exports."""

    if not rows:
        return None
    try:
        import matplotlib

        matplotlib.use("Agg", force=True)
        import matplotlib.pyplot as plt
        from matplotlib.patches import Patch
    except Exception:
        return None

    labels = [f"{_text(row.get('scan_type'), 'Scan')} #{index}" for index, row in enumerate(rows, 1)]
    values = [_percent(row.get("confidence")) for row in rows]
    predictions = [_text(row.get("prediction"), "Unknown") for row in rows]
    colors = [_prediction_color(prediction) for prediction in predictions]

    fig_width = max(7.2, min(12.0, 0.72 * len(rows) + 4.2))
    fig, ax = plt.subplots(figsize=(fig_width, 3.4), dpi=160)
    fig.patch.set_facecolor("#FFFFFF")
    ax.set_facecolor("#FFFFFF")
    ax.bar(range(len(rows)), values, color=colors, width=0.58)
    ax.set_title("Selected evidence confidence overview", fontsize=12, fontweight="bold", color="#0F172A", pad=12)
    ax.set_ylabel("Confidence (%)", fontsize=9, color="#334155")
    ax.set_xlabel("Scan evidence", fontsize=9, color="#334155")
    ax.set_ylim(0, 100)
    ax.set_xticks(range(len(rows)))
    ax.set_xticklabels(labels, rotation=32, ha="right", fontsize=7.5, color="#334155")
    ax.tick_params(axis="y", labelsize=8, colors="#334155")
    ax.grid(axis="y", color="#CBD5E1", linewidth=0.7, alpha=0.7)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.spines["left"].set_color("#CBD5E1")
    ax.spines["bottom"].set_color("#CBD5E1")

    legend_items: list[Patch] = []
    seen: set[str] = set()
    for prediction in predictions:
        if prediction in seen:
            continue
        seen.add(prediction)
        legend_items.append(Patch(facecolor=_prediction_color(prediction), label=prediction))
    if legend_items:
        ax.legend(handles=legend_items, loc="upper right", fontsize=7.5, frameon=False)

    buffer = io.BytesIO()
    fig.tight_layout()
    fig.savefig(buffer, format="png", bbox_inches="tight")
    plt.close(fig)
    return buffer.getvalue()


def _confidence_chart_lines(rows: list[dict[str, object]]) -> list[str]:
    lines = ["Confidence Overview"]
    for index, row in enumerate(rows, 1):
        lines.append(
            f"- {index}. {_text(row.get('scan_type'))}: {_text(row.get('prediction'))} "
            f"({_percent(row.get('confidence')):.1f}%)"
        )
    return lines


def _filename(extension: str) -> str:
    stamp = now_for_app().strftime("%Y%m%d_%H%M%S")
    return f"AIFDS_Report_{stamp}.{extension.lower()}"


def build_preview(rows: list[dict[str, object]], report_note: str, sections: dict[str, bool]) -> str:
    counts = _risk_counts(rows)
    lines = [
        "AI-based Spam and Caller Fraud Detection System",
        "AI Analysis Evidence Report",
        f"Generated: {formatted_now()}",
        f"Records included: {len(rows)}",
        "",
    ]
    if sections.get("summary", True):
        lines.extend(
            [
                "Executive Summary",
                f"- High risk: {counts.get('High risk', 0)}",
                f"- Needs review: {counts.get('Needs review', 0)}",
                f"- Lower risk: {counts.get('Lower risk', 0)}",
                "",
            ]
        )
    if sections.get("evidence", True):
        lines.append("Evidence Included")
        for index, row in enumerate(rows, 1):
            lines.append(
                f"{index}. {_date_text(row.get('scanned_at'))} | "
                f"{_text(row.get('scan_type'))} | {_text(row.get('prediction'))} | "
                f"{_percent(row.get('confidence')):.1f}%"
            )
        lines.append("")
    if sections.get("risk", True):
        lines.extend(_confidence_chart_lines(rows))
        lines.append("")
    if sections.get("recommendations", True):
        lines.extend(["Recommendations", report_note.strip() or DEFAULT_RECOMMENDATION, ""])
    if sections.get("appendix", True):
        lines.extend(
            [
                "Scope",
                "This report is an educational capstone prototype output. It supports awareness and review, not legal or forensic proof.",
            ]
        )
    return "\n".join(lines)


def build_txt(rows: list[dict[str, object]], report_note: str, sections: dict[str, bool]) -> tuple[bytes, str]:
    body = build_preview(rows, report_note, sections)
    if sections.get("explanations", True):
        detail_lines = ["", "Individual Scan Details"]
        for index, row in enumerate(rows, 1):
            detail_lines.extend(
                [
                    f"Scan #{index}",
                    f"Time: {_date_text(row.get('scanned_at'))}",
                    f"Type: {_text(row.get('scan_type'))}",
                    f"Prediction: {_text(row.get('prediction'))}",
                    f"Confidence: {_percent(row.get('confidence')):.1f}%",
                    f"Model: {_text(row.get('model_name'))}",
                    f"Source: {_text(row.get('source_name'))}",
                    f"Flags: {', '.join(_flags(row.get('flags'))) or '-'}",
                    f"Explanation: {_text(row.get('explanation'))}",
                    f"Preview: {_text(row.get('preview'))}",
                    "",
                ]
            )
        body += "\n".join(detail_lines)
    return body.encode("utf-8"), _filename("txt")


def build_pdf(rows: list[dict[str, object]], report_note: str, sections: dict[str, bool]) -> tuple[bytes, str]:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.platypus import Image, PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=1.6 * cm,
        leftMargin=1.6 * cm,
        topMargin=1.5 * cm,
        bottomMargin=1.5 * cm,
        title="AI-FDS Analysis Evidence Report",
    )
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="Muted",
            parent=styles["BodyText"],
            fontSize=9,
            leading=12,
            textColor=colors.HexColor("#475569"),
        )
    )
    styles.add(
        ParagraphStyle(
            name="SmallMono",
            parent=styles["BodyText"],
            fontName="Courier",
            fontSize=8,
            leading=10,
            textColor=colors.HexColor("#334155"),
        )
    )
    story = [
        Paragraph("AI-based Spam and Caller Fraud Detection System", styles["Title"]),
        Paragraph("AI Analysis Evidence Report", styles["Heading2"]),
        Paragraph(html.escape(f"Generated: {formatted_now()}"), styles["Muted"]),
        Paragraph("Educational capstone prototype. Not legal or forensic proof.", styles["Muted"]),
        Spacer(1, 12),
    ]

    if sections.get("summary", True):
        counts = _risk_counts(rows)
        story.append(Paragraph("1. Executive Summary", styles["Heading2"]))
        summary_rows = [
            ["Records included", str(len(rows))],
            ["High risk", str(counts.get("High risk", 0))],
            ["Needs review", str(counts.get("Needs review", 0))],
            ["Lower risk", str(counts.get("Lower risk", 0))],
        ]
        table = Table(summary_rows, colWidths=[6 * cm, 9 * cm])
        table.setStyle(
            TableStyle(
                [
                    ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#CBD5E1")),
                    ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#E0F2FE")),
                    ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                    ("TOPPADDING", (0, 0), (-1, -1), 6),
                ]
            )
        )
        story.extend([table, Spacer(1, 12)])

    if sections.get("evidence", True):
        story.append(Paragraph("2. Scan Evidence Table", styles["Heading2"]))
        table_rows = [["#", "Time", "Type", "Prediction", "Conf.", "Model"]]
        for index, row in enumerate(rows, 1):
            table_rows.append(
                [
                    str(index),
                    _date_text(row.get("scanned_at"))[:16],
                    _text(row.get("scan_type")),
                    _text(row.get("prediction")),
                    f"{_percent(row.get('confidence')):.0f}%",
                    _text(row.get("model_name"))[:28],
                ]
            )
        table = Table(table_rows, repeatRows=1, colWidths=[0.8 * cm, 3.2 * cm, 2.3 * cm, 3.5 * cm, 1.5 * cm, 4 * cm])
        table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0F766E")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#CBD5E1")),
                    ("FONTSIZE", (0, 0), (-1, -1), 8),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                    ("TOPPADDING", (0, 0), (-1, -1), 5),
                ]
            )
        )
        story.extend([table, Spacer(1, 12)])

    if sections.get("explanations", True):
        story.append(Paragraph("3. Individual Scan Detail", styles["Heading2"]))
        for index, row in enumerate(rows, 1):
            story.append(Paragraph(html.escape(f"Scan #{index}: {_text(row.get('scan_type'))}"), styles["Heading3"]))
            detail = (
                f"Prediction: {_text(row.get('prediction'))} | "
                f"Confidence: {_percent(row.get('confidence')):.1f}% | "
                f"Model: {_text(row.get('model_name'))} | "
                f"Source: {_text(row.get('source_name'))}"
            )
            story.append(Paragraph(html.escape(detail), styles["Muted"]))
            flags = ", ".join(_flags(row.get("flags")))
            if flags:
                story.append(Paragraph(html.escape(f"Flags: {flags}"), styles["Muted"]))
            explanation = _text(row.get("explanation"), "")
            if explanation:
                story.append(Paragraph(html.escape(explanation), styles["Muted"]))
            preview = _text(row.get("preview"), "")
            if preview:
                story.append(Paragraph(html.escape(preview[:350]), styles["SmallMono"]))
            story.append(Spacer(1, 8))

    if sections.get("risk", True):
        story.append(Paragraph("4. Confidence Overview and Risk Interpretation", styles["Heading2"]))
        chart = _confidence_chart_png(rows)
        if chart:
            story.append(Image(io.BytesIO(chart), width=15.2 * cm, height=6.2 * cm))
            story.append(Spacer(1, 8))
        story.append(
            Paragraph(
                "High confidence does not prove fraud by itself. It means the prototype found patterns similar to its training or rule examples. "
                "Human review is still required before taking action.",
                styles["Muted"],
            )
        )
        story.append(Spacer(1, 12))

    if sections.get("recommendations", True):
        story.append(Paragraph("5. Recommendations", styles["Heading2"]))
        for item in [
            "Pause before responding to urgent financial, account, or identity requests.",
            "Verify the sender or caller through a separate official channel.",
            "Do not share OTP codes, passwords, TAC codes, banking details, or identity numbers.",
            "Preserve screenshots, emails, transcripts, audio filenames, and timestamps for review.",
            "Escalate suspicious campus-related messages to the appropriate university support or security team.",
        ]:
            story.append(Paragraph(html.escape(f"- {item}"), styles["Muted"]))
        story.append(Paragraph(html.escape(report_note.strip() or DEFAULT_RECOMMENDATION), styles["Muted"]))
        story.append(Spacer(1, 12))

    if sections.get("appendix", True):
        story.append(PageBreak())
        story.append(Paragraph("6. Appendix and Scope", styles["Heading2"]))
        appendix = [
            ["System", "AI-FDS Capstone Prototype"],
            ["Included models", "TF-IDF + Naive Bayes / Decision Tree, MFCC + SVM, rule-based educational fallbacks"],
            ["Data status", "Synthetic examples may be active until official datasets and trained model artifacts are inserted."],
            ["Scope", "Educational scam awareness support. Not enterprise security, telecom verification, or legal evidence."],
        ]
        table = Table(appendix, colWidths=[4 * cm, 11 * cm])
        table.setStyle(
            TableStyle(
                [
                    ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#CBD5E1")),
                    ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#F1F5F9")),
                    ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
                    ("FONTSIZE", (0, 0), (-1, -1), 8.5),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                    ("TOPPADDING", (0, 0), (-1, -1), 6),
                ]
            )
        )
        story.append(table)

    doc.build(story)
    return buffer.getvalue(), _filename("pdf")


def build_docx(rows: list[dict[str, object]], report_note: str, sections: dict[str, bool]) -> tuple[bytes, str]:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

    document = Document()
    title = document.add_heading("AI-based Spam and Caller Fraud Detection System", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = document.add_paragraph("AI Analysis Evidence Report")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta = document.add_paragraph(f"Generated: {formatted_now()}\nRecords included: {len(rows)}\nEducational capstone prototype. Not legal or forensic proof.")
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def heading(text: str, level: int = 1) -> None:
        paragraph = document.add_heading(text, level=level)
        if paragraph.runs:
            paragraph.runs[0].font.color.rgb = RGBColor(15, 118, 110)

    def body(text: str) -> None:
        paragraph = document.add_paragraph(text)
        for run in paragraph.runs:
            run.font.size = Pt(10)

    if sections.get("summary", True):
        heading("1. Executive Summary")
        counts = _risk_counts(rows)
        body(
            f"Total records: {len(rows)} | High risk: {counts.get('High risk', 0)} | "
            f"Needs review: {counts.get('Needs review', 0)} | Lower risk: {counts.get('Lower risk', 0)}."
        )

    if sections.get("evidence", True):
        heading("2. Scan Evidence Table")
        table = document.add_table(rows=1, cols=6)
        table.style = "Table Grid"
        for index, label in enumerate(["#", "Timestamp", "Type", "Prediction", "Conf.", "Model"]):
            table.rows[0].cells[index].text = label
        for index, row in enumerate(rows, 1):
            cells = table.add_row().cells
            values = [
                str(index),
                _date_text(row.get("scanned_at"))[:16],
                _text(row.get("scan_type")),
                _text(row.get("prediction")),
                f"{_percent(row.get('confidence')):.0f}%",
                _text(row.get("model_name"))[:35],
            ]
            for cell_index, value in enumerate(values):
                cells[cell_index].text = value

    if sections.get("explanations", True):
        heading("3. Individual Scan Detail")
        for index, row in enumerate(rows, 1):
            heading(f"Scan #{index}: {_text(row.get('scan_type'))}", level=2)
            body(
                f"Time: {_date_text(row.get('scanned_at'))}\n"
                f"Prediction: {_text(row.get('prediction'))} ({_percent(row.get('confidence')):.1f}%)\n"
                f"Model: {_text(row.get('model_name'))}\n"
                f"Source: {_text(row.get('source_name'))}"
            )
            flags = ", ".join(_flags(row.get("flags")))
            if flags:
                body(f"Flags: {flags}")
            explanation = _text(row.get("explanation"), "")
            if explanation:
                body(f"Explanation: {explanation}")
            preview = _text(row.get("preview"), "")
            if preview:
                body(f"Evidence preview: {preview[:350]}")

    if sections.get("risk", True):
        heading("4. Confidence Overview and Risk Interpretation")
        chart = _confidence_chart_png(rows)
        if chart:
            document.add_picture(io.BytesIO(chart), width=Inches(6.4))
        body(
            "Model confidence is a prototype decision signal, not absolute proof. "
            "High-risk or unclear items should be reviewed by a person through official channels."
        )

    if sections.get("recommendations", True):
        heading("5. Recommendations")
        for item in [
            "Pause before responding to urgent financial, account, or identity requests.",
            "Verify the sender or caller through a separate official channel.",
            "Do not share OTP codes, passwords, TAC codes, banking details, or identity numbers.",
            "Preserve screenshots, emails, transcripts, audio filenames, and timestamps for review.",
            "Escalate suspicious campus-related messages to the appropriate university support or security team.",
        ]:
            document.add_paragraph(item, style="List Bullet")
        body(report_note.strip() or DEFAULT_RECOMMENDATION)

    if sections.get("appendix", True):
        document.add_page_break()
        heading("6. Appendix and Scope")
        body(
            "System: AI-FDS Capstone Prototype\n"
            "Included models: TF-IDF + Naive Bayes / Decision Tree, MFCC + SVM, rule-based educational fallbacks\n"
            "Data status: Synthetic examples may be active until official datasets and trained model artifacts are inserted.\n"
            "Scope: Educational scam awareness support. Not enterprise security, telecom verification, or legal evidence."
        )

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue(), _filename("docx")


def build_report(
    report_format: Literal["TXT", "PDF", "DOCX"],
    rows: list[dict[str, object]],
    report_note: str,
    sections: dict[str, bool],
) -> tuple[bytes, str, str]:
    """Build a report and return bytes, filename, and MIME type."""

    if report_format == "TXT":
        payload, filename = build_txt(rows, report_note, sections)
        return payload, filename, "text/plain"
    if report_format == "PDF":
        payload, filename = build_pdf(rows, report_note, sections)
        return payload, filename, "application/pdf"
    if report_format == "DOCX":
        payload, filename = build_docx(rows, report_note, sections)
        return payload, filename, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    raise ValueError(f"Unsupported report format: {report_format}")
