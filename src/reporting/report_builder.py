"""Report builders for the AI-FDS report generator page."""

from __future__ import annotations

import html
import hashlib
import io
import json
import os
import re
import shutil
import subprocess
import sys
import tempfile
from collections import Counter
from pathlib import Path
from typing import Literal

from src.reporting.evidence_snapshot import (
    ACTION_STATUSES,
    IMMEDIATE_ACTION,
    NO_IMMEDIATE_ACTION,
    REVIEW_REQUIRED,
    decode_json_object,
    derive_action_status,
    remediation_plan,
)
from src.utils.time_utils import formatted_now, now_for_app


THEME = {
    "navy": "#0B1220",
    "surface": "#111827",
    "blue": "#2563EB",
    "cyan": "#0891B2",
    "violet": "#7C3AED",
    "green": "#059669",
    "orange": "#F97316",
    "red": "#DC2626",
    "muted": "#475569",
    "border": "#CBD5E1",
    "soft_blue": "#DBEAFE",
    "soft_violet": "#EDE9FE",
    "soft_orange": "#FFEDD5",
    "soft_green": "#DCFCE7",
}


DEFAULT_RECOMMENDATION = (
    "Verify suspicious contact through official channels, do not share OTP/passwords, "
    "preserve the evidence, and report urgent financial or identity-related requests "
    "to the relevant campus or organisation support team."
)

DEFAULT_SECTIONS = {
    "summary": True,
    "evidence": True,
    "explanations": True,
    "risk": True,
    "recommendations": True,
    "appendix": True,
}

STUDENT_PROFILE = "Student Brief"
TECHNICAL_PROFILE = "Technical Report"
REPORT_PROFILES = [STUDENT_PROFILE, TECHNICAL_PROFILE]
REPORT_SCHEMA_VERSION = "student-brief-evidence-artifacts-v2"

DOCUMENTATION_REFERENCES = [
    (
        "EC-Council CEH - Social Engineering, Phishing, AI Impersonation, and Countermeasures",
        "https://www.eccouncil.org/train-certify/certified-ethical-hacker-ceh/",
    ),
    (
        "NIST SP 800-61 Rev. 3 - Incident Response Recommendations and Considerations",
        "https://csrc.nist.gov/pubs/sp/800/61/r3/final",
    ),
    (
        "SWGDE Requirements for Report Writing in Digital and Multimedia Forensics",
        "https://www.swgde.org/documents/published-complete-listing/18-q-002-swgde-requirements-for-report-writing-in-digital-and-multimedia-forensics/",
    ),
    (
        "SWGDE Best Practices for Digital Evidence Collection",
        "https://www.swgde.org/documents/published-complete-listing/18-f-002-best-practices-for-digital-evidence-collection/",
    ),
    (
        "NISTIR 8312 - Four Principles of Explainable Artificial Intelligence",
        "https://www.nist.gov/publications/four-principles-explainable-artificial-intelligence",
    ),
    (
        "CISA Secure Our World - Recognize and Report Phishing",
        "https://www.cisa.gov/secure-our-world",
    ),
    (
        "FTC Consumer Advice - AI Voice Cloning and Family Emergency Scams",
        "https://consumer.ftc.gov/articles/scammers-use-fake-emergencies-steal-your-money",
    ),
    (
        "FCC Consumer Complaints - Unwanted Calls, Texts, and Caller ID Spoofing",
        "https://consumercomplaints.fcc.gov/hc/en-us/articles/115002234203-Unwanted-Calls-Texts-Phone",
    ),
]


def _text(value: object, fallback: str = "-") -> str:
    if value is None:
        return fallback
    value_text = str(value).strip()
    return value_text or fallback


def _percent(value: object) -> float:
    try:
        return max(0.0, min(100.0, float(value)))
    except (TypeError, ValueError):
        return 0.0


def _date_text(value: object) -> str:
    text = _text(value)
    return text.replace("T", " ")[:19]


def _context_text(
    context: dict[str, object] | None,
    key: str,
    fallback: str,
) -> str:
    if not context:
        return fallback
    return _text(context.get(key), fallback)


def _flags(value: object) -> list[str]:
    if value is None:
        return []
    if isinstance(value, str):
        try:
            return _flags(json.loads(value))
        except json.JSONDecodeError:
            return [value] if value.strip() else []
    if isinstance(value, list):
        return [_text(item, "") for item in value if _text(item, "")]
    return [_text(value)]


def _evidence_family(row: dict[str, object]) -> str:
    scan_type = _text(row.get("scan_type"), "").strip().casefold()
    exact_families = {
        "email": "Email",
        "transcript": "Transcript",
        "audio": "Audio",
        "phone": "Phone",
    }
    if scan_type in exact_families:
        return exact_families[scan_type]

    text = " ".join(
        [
            _text(row.get("scan_type"), ""),
            _text(row.get("source_name"), ""),
            _text(row.get("model_name"), ""),
            _text(row.get("prediction"), ""),
        ]
    ).lower()
    if any(term in text for term in ("phone", "caller", "veriphone", "carrier")):
        return "Phone"
    if any(term in text for term in ("audio", "voice", "deepfake", "mfcc", "speaker", "recording")):
        return "Audio"
    if any(term in text for term in ("transcript", "call", "meeting", "whisper")):
        return "Transcript"
    if any(term in text for term in ("email", "message", "mail", "sms")):
        return "Email"
    return _text(row.get("scan_type"), "Evidence")


def _is_unknown_result(row: dict[str, object]) -> bool:
    prediction = _text(row.get("native_prediction") or row.get("prediction"), "").lower()
    if any(term in prediction for term in ("unknown", "unavailable", "not found", "reputation unknown")):
        return True
    return _evidence_family(row) == "Phone" and _percent(row.get("confidence")) <= 0


def _score_available(row: dict[str, object]) -> bool:
    explicit = row.get("score_available")
    if isinstance(explicit, str):
        return explicit.strip().casefold() in {"true", "1", "yes"}
    if explicit is not None:
        return bool(explicit)
    return row.get("concern_score") is not None


def _score_text(row: dict[str, object]) -> str:
    if not _score_available(row):
        return "N/A"
    score = row.get("concern_score")
    if score is None and _evidence_family(row) == "Transcript":
        score = row.get("confidence")
    return f"{_percent(score):.1f}%"


def _risk_bucket(row: dict[str, object]) -> str:
    stored = _text(row.get("action_status"), "")
    if stored in ACTION_STATUSES:
        return stored
    return derive_action_status(
        native_prediction=row.get("native_prediction") or row.get("prediction"),
        evidence_type=row.get("scan_type"),
        concern_score=row.get("concern_score"),
        score_available=_score_available(row),
    )


def _risk_counts(rows: list[dict[str, object]]) -> Counter:
    return Counter(_risk_bucket(row) for row in rows)


def _prediction_color(prediction: object) -> str:
    text = _text(prediction, "")
    if text == NO_IMMEDIATE_ACTION:
        return THEME["green"]
    if text == IMMEDIATE_ACTION:
        return THEME["red"]
    return THEME["orange"]


def _evidence_counts(rows: list[dict[str, object]]) -> Counter:
    return Counter(_evidence_family(row) for row in rows)


def _short_source(row: dict[str, object]) -> str:
    source = _text(row.get("source_name"), "")
    if source:
        return source[:70]
    family = _evidence_family(row)
    if family == "Phone":
        return "Veriphone.io lookup"
    if family == "Audio":
        return "Recorded or uploaded audio"
    if family == "Transcript":
        return "Uploaded or pasted transcript"
    if family == "Email":
        return "Uploaded or pasted message"
    return "-"


def _engine_text(row: dict[str, object]) -> str:
    engine = _text(row.get("model_name"), "")
    if engine:
        return engine
    if _evidence_family(row) == "Phone":
        return "Veriphone.io + PenipuMY"
    return "-"


def _bundle(row: dict[str, object]) -> dict[str, object]:
    return decode_json_object(row.get("evidence_bundle"))


def _provenance(row: dict[str, object]) -> dict[str, object]:
    return decode_json_object(row.get("provenance"))


def _native_prediction(row: dict[str, object]) -> str:
    return _text(row.get("native_prediction") or row.get("prediction"), "Unknown")


def _browser_path() -> str:
    configured = os.environ.get("BROWSER_PATH", "").strip()
    if configured and Path(configured).exists():
        return configured
    candidates = [
        shutil.which("google-chrome"),
        shutil.which("chromium"),
        shutil.which("chromium-browser"),
        shutil.which("msedge"),
        r"C:\Program Files\Google\Chrome\Application\chrome.exe",
        r"C:\Program Files (x86)\Google\Chrome\Application\chrome.exe",
        r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe",
    ]
    return next(
        (
            str(candidate)
            for candidate in candidates
            if candidate and Path(str(candidate)).exists()
        ),
        "",
    )


def _artifact_figure_key(artifact: dict[str, object]) -> str:
    saved = _text(artifact.get("figure_sha256"), "")
    if saved:
        return saved
    figure_json = _text(artifact.get("figure_json"), "")
    return hashlib.sha256(figure_json.encode("utf-8")).hexdigest() if figure_json else ""


def _ordered_chart_artifacts(rows: list[dict[str, object]]) -> list[dict[str, object]]:
    artifacts: list[dict[str, object]] = []
    seen: set[str] = set()
    for row in rows:
        bundle_artifacts = _bundle(row).get("artifacts", [])
        if not isinstance(bundle_artifacts, list):
            continue
        for artifact in bundle_artifacts:
            if not isinstance(artifact, dict) or artifact.get("kind") != "chart":
                continue
            key = _artifact_figure_key(artifact)
            if key and key not in seen and artifact.get("figure_json"):
                seen.add(key)
                artifacts.append(artifact)
    return artifacts


def _render_artifact_pngs(rows: list[dict[str, object]]) -> dict[str, bytes]:
    """Render all saved Plotly charts once, with a killable hard timeout."""

    artifacts = _ordered_chart_artifacts(rows)
    if not artifacts:
        return {}

    timeout_seconds = 60.0
    try:
        timeout_seconds = max(
            10.0,
            min(180.0, float(os.environ.get("AIFDS_CHART_RENDER_TIMEOUT", "60"))),
        )
    except ValueError:
        pass

    root = Path(__file__).resolve().parents[2]
    temp_root = root / "tmp"
    temp_root.mkdir(parents=True, exist_ok=True)
    env = os.environ.copy()
    browser_path = _browser_path()
    if browser_path:
        env["BROWSER_PATH"] = browser_path

    creation_flags = getattr(subprocess, "CREATE_NO_WINDOW", 0)
    try:
        with tempfile.TemporaryDirectory(
            prefix="aifds_chart_render_",
            dir=temp_root,
        ) as temp_dir:
            temp_path = Path(temp_dir)
            runtime_temp = temp_path / "runtime"
            runtime_temp.mkdir(exist_ok=True)
            env["TEMP"] = str(runtime_temp)
            env["TMP"] = str(runtime_temp)
            env["TMPDIR"] = str(runtime_temp)
            output_dir = temp_path / "images"
            output_dir.mkdir(exist_ok=True)
            manifest = []
            for index, artifact in enumerate(artifacts):
                key = _artifact_figure_key(artifact)
                manifest.append(
                    {
                        "figure_json": artifact.get("figure_json", ""),
                        "filename": f"{index:04d}-{key}.png",
                        "key": key,
                    }
                )
            manifest_path = temp_path / "manifest.json"
            manifest_path.write_text(
                json.dumps(manifest, ensure_ascii=True),
                encoding="utf-8",
            )
            subprocess.run(
                [
                    sys.executable,
                    "-m",
                    "src.reporting.chart_renderer",
                    str(manifest_path),
                    str(output_dir),
                ],
                cwd=root,
                env=env,
                capture_output=True,
                text=True,
                timeout=timeout_seconds,
                check=True,
                creationflags=creation_flags,
            )
            rendered: dict[str, bytes] = {}
            for entry in manifest:
                image_path = output_dir / str(entry["filename"])
                if image_path.exists():
                    rendered[str(entry["key"])] = image_path.read_bytes()
            return rendered
    except (OSError, subprocess.SubprocessError, ValueError):
        return {}


def _artifact_png(
    artifact: dict[str, object],
    image_cache: dict[str, bytes],
) -> bytes | None:
    key = _artifact_figure_key(artifact)
    return image_cache.get(key) if key else None


def _artifact_rows(artifact: dict[str, object]) -> list[dict[str, object]]:
    rows = artifact.get("rows")
    if isinstance(rows, list):
        return [row for row in rows if isinstance(row, dict)]
    data = artifact.get("data")
    if isinstance(data, list):
        return [row for row in data if isinstance(row, dict)]
    if isinstance(data, dict):
        return [{"Field": key, "Value": value} for key, value in data.items()]
    return []


def _artifact_columns(artifact: dict[str, object], rows: list[dict[str, object]]) -> list[str]:
    columns = artifact.get("columns")
    if isinstance(columns, list) and columns:
        return [str(column) for column in columns]
    output: list[str] = []
    for row in rows:
        for key in row:
            if str(key) not in output:
                output.append(str(key))
    return output


def _chunked_columns(columns: list[str], size: int = 5) -> list[list[str]]:
    return [columns[index : index + size] for index in range(0, len(columns), size)]


def _value_text(value: object, limit: int = 220) -> str:
    if isinstance(value, (dict, list)):
        text = json.dumps(value, ensure_ascii=True, default=str)
    else:
        text = _text(value, "")
    return text if len(text) <= limit else text[: limit - 3].rstrip() + "..."


def _bundle_summary_rows(row: dict[str, object]) -> list[tuple[str, str]]:
    bundle = _bundle(row)
    dashboard = bundle.get("dashboard_summary", {})
    values = [
        ("Native verdict", _native_prediction(row)),
        ("Action status", _risk_bucket(row)),
        (_text(row.get("score_label"), "Concern score"), _score_text(row)),
    ]
    if isinstance(dashboard, dict):
        for key, value in dashboard.items():
            label = str(key).replace("_", " ").strip().title()
            text = _value_text(value)
            if text and (label, text) not in values:
                values.append((label, text))
    return values


def _provenance_rows(row: dict[str, object]) -> list[tuple[str, str]]:
    provenance = _provenance(row)
    if not provenance:
        return [("Provenance", "Legacy record: no immutable provenance snapshot was saved.")]
    keys = (
        "captured_at",
        "source_name",
        "source_kind",
        "sha256",
        "hash_scope",
        "integrity_note",
        "normalization",
        "provider_coverage",
    )
    labels = {
        "sha256": "Input SHA-256",
        "captured_at": "Captured At",
    }
    return [
        (labels.get(key, key.replace("_", " ").title()), _value_text(provenance.get(key), 500))
        for key in keys
        if provenance.get(key) not in (None, "", [], {})
    ]


def _xai_rows(row: dict[str, object]) -> list[dict[str, object]]:
    xai = _bundle(row).get("xai", {})
    if not isinstance(xai, dict):
        return []
    factors = xai.get("factors", [])
    return [item for item in factors if isinstance(item, dict)] if isinstance(factors, list) else []


def _remediation_items(
    rows: list[dict[str, object]],
    incident_context: dict[str, object] | None = None,
) -> list[dict[str, str]]:
    output: list[dict[str, str]] = []
    for row in rows:
        bundle = _bundle(row)
        saved = bundle.get("remediation", [])
        candidates = [item for item in saved if isinstance(item, dict)] if isinstance(saved, list) else []
        candidates.extend(
            remediation_plan(
                evidence_type=_evidence_family(row),
                action_status=_risk_bucket(row),
                findings=_flags(row.get("flags")),
                incident_context=incident_context,
            )
        )
        for item in candidates:
            normalized = {
                "priority": _text(item.get("priority"), "Review"),
                "trigger": _text(item.get("trigger"), "Selected evidence"),
                "action": _text(item.get("action"), "Verify through an independent official channel."),
                "reason": _text(item.get("reason"), "Reduces the risk of relying on unverified contact."),
            }
            if normalized not in output:
                output.append(normalized)
    priority_order = {"Immediate": 0, "Prompt review": 1, "Verify": 2, "Preserve": 3, "Monitor": 4}
    return sorted(output, key=lambda item: priority_order.get(item["priority"], 5))


def _artifact_lines(row: dict[str, object]) -> list[str]:
    bundle = _bundle(row)
    artifacts = bundle.get("artifacts", [])
    if not isinstance(artifacts, list) or not artifacts:
        return ["Dashboard artifacts", "- Legacy record: no dashboard snapshot was saved."]

    lines = ["Dashboard artifacts (saved display order)"]
    for index, artifact in enumerate(artifacts, 1):
        if not isinstance(artifact, dict):
            continue
        kind = _text(artifact.get("kind"), "artifact").title()
        title = _text(artifact.get("title"), f"Artifact {index}")
        lines.append(f"{index}. [{kind}] {title}")
        description = _text(artifact.get("description"), "")
        if description:
            lines.append(f"   Description: {description}")
        if artifact.get("figure_sha256"):
            lines.append(f"   Figure snapshot SHA-256: {artifact['figure_sha256']}")
        if artifact.get("body"):
            lines.append(f"   Content: {_text(artifact.get('body'), '')}")
        rows = _artifact_rows(artifact)
        columns = _artifact_columns(artifact, rows)
        if rows and columns:
            lines.append("   Data:")
            for row_data in rows:
                values = " | ".join(
                    f"{column}: {_value_text(row_data.get(column), 500)}"
                    for column in columns
                )
                lines.append(f"   - {values}")
        elif kind == "Chart" and not artifact.get("figure_json"):
            lines.append("   Availability: chart data was unavailable when the snapshot was captured.")
    return lines


def _indicator_label(value: str) -> str:
    text = " ".join(value.replace("_", " ").split()).strip(" .:-")
    if not text:
        return "Evidence indicator"
    if len(text) > 54:
        text = text[:51].rstrip() + "..."
    return text[:1].upper() + text[1:]


def _indicator_counts(rows: list[dict[str, object]]) -> Counter:
    counts: Counter = Counter()
    for row in rows:
        for flag in _flags(row.get("flags")):
            counts[_indicator_label(flag)] += 1
        if _is_unknown_result(row) and _evidence_family(row) == "Phone":
            counts["Phone ownership unconfirmed"] += 1
    return counts


def _combined_findings(rows: list[dict[str, object]]) -> list[dict[str, str]]:
    evidence_by_indicator: dict[str, set[str]] = {}
    severity_by_indicator: dict[str, str] = {}
    for index, row in enumerate(rows, 1):
        family = _evidence_family(row)
        bucket = _risk_bucket(row)
        labels = [_indicator_label(flag) for flag in _flags(row.get("flags"))]
        if not labels and _is_unknown_result(row) and family == "Phone":
            labels = ["Phone ownership unconfirmed"]
        for label in labels:
            evidence_by_indicator.setdefault(label, set()).add(f"{family} #{index}")
            current = severity_by_indicator.get(label, "Informative")
            severity_by_indicator[label] = _stronger_severity(current, _severity_from_bucket(bucket))

    rows_out = []
    for label, evidence_set in sorted(evidence_by_indicator.items(), key=lambda item: (-len(item[1]), item[0]))[:10]:
        rows_out.append(
            {
                "Finding": label,
                "Evidence involved": ", ".join(sorted(evidence_set)),
                "Severity": severity_by_indicator.get(label, "Informative"),
            }
        )
    return rows_out


def _severity_from_bucket(bucket: str) -> str:
    if bucket == IMMEDIATE_ACTION:
        return "High"
    if bucket == REVIEW_REQUIRED:
        return "Review"
    return "Informative"


def _stronger_severity(left: str, right: str) -> str:
    order = {"Informative": 0, "Unknown": 1, "Review": 2, "High": 3}
    return left if order.get(left, 0) >= order.get(right, 0) else right


def _summary_rows(rows: list[dict[str, object]]) -> list[tuple[str, str]]:
    counts = _risk_counts(rows)
    evidence_counts = _evidence_counts(rows)
    return [
        ("Evidence selected", str(len(rows))),
        ("Email evidence", str(evidence_counts.get("Email", 0))),
        ("Transcript evidence", str(evidence_counts.get("Transcript", 0))),
        ("Audio evidence", str(evidence_counts.get("Audio", 0))),
        ("Phone evidence", str(evidence_counts.get("Phone", 0))),
        (IMMEDIATE_ACTION, str(counts.get(IMMEDIATE_ACTION, 0))),
        (REVIEW_REQUIRED, str(counts.get(REVIEW_REQUIRED, 0))),
        (NO_IMMEDIATE_ACTION, str(counts.get(NO_IMMEDIATE_ACTION, 0))),
    ]


def _confidence_chart_png(rows: list[dict[str, object]]) -> bytes | None:
    """Render available classification scores as PNG for PDF/DOCX exports."""

    score_rows = [(index, row) for index, row in enumerate(rows, 1) if _score_available(row)]
    if not score_rows:
        return None
    try:
        import matplotlib

        matplotlib.use("Agg", force=True)
        import matplotlib.pyplot as plt
        from matplotlib.patches import Patch
    except Exception:
        return None

    labels = [f"{_evidence_family(row)} #{index}" for index, row in score_rows]
    values = [
        _percent(
            row.get("concern_score")
            if row.get("concern_score") is not None
            else row.get("confidence")
        )
        for _, row in score_rows
    ]
    predictions = [_risk_bucket(row) for _, row in score_rows]
    colors = [_prediction_color(prediction) for prediction in predictions]

    fig_height = max(3.2, min(7.5, 0.45 * len(score_rows) + 2.2))
    fig, ax = plt.subplots(figsize=(8.2, fig_height), dpi=160)
    fig.patch.set_facecolor("#FFFFFF")
    ax.set_facecolor("#FFFFFF")
    ax.barh(range(len(score_rows)), values, color=colors, height=0.56)
    ax.set_title("Normalized Action Overview", fontsize=12, fontweight="bold", color=THEME["navy"], pad=12)
    ax.set_xlabel("Source-native concern score where available (%)", fontsize=9, color="#334155")
    ax.set_xlim(0, 100)
    ax.set_yticks(range(len(score_rows)))
    ax.set_yticklabels(labels, fontsize=8.5, color="#334155")
    ax.invert_yaxis()
    ax.tick_params(axis="x", labelsize=8, colors="#334155")
    ax.grid(axis="x", color=THEME["border"], linewidth=0.7, alpha=0.7)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.spines["left"].set_color(THEME["border"])
    ax.spines["bottom"].set_color(THEME["border"])

    legend_items: list[Patch] = []
    seen: set[str] = set()
    for prediction in predictions:
        if prediction in seen:
            continue
        seen.add(prediction)
        legend_items.append(Patch(facecolor=_prediction_color(prediction), label=prediction))
    if legend_items:
        ax.legend(handles=legend_items, loc="upper right", fontsize=7.5, frameon=False)

    buffer = io.BytesIO()
    fig.tight_layout()
    fig.savefig(buffer, format="png", bbox_inches="tight")
    plt.close(fig)
    return buffer.getvalue()


def _confidence_chart_lines(rows: list[dict[str, object]]) -> list[str]:
    lines = ["Normalized Action Overview"]
    for index, row in enumerate(rows, 1):
        lines.append(
            f"- {_evidence_family(row)} #{index}: {_risk_bucket(row)} | "
            f"Native verdict: {_text(row.get('native_prediction') or row.get('prediction'))} | "
            f"{_text(row.get('score_label'), 'Concern score')}: {_score_text(row)}"
        )
    return lines


def _count_chart_png(title: str, counts: Counter, color: str) -> bytes | None:
    if not counts:
        return None
    try:
        import matplotlib

        matplotlib.use("Agg", force=True)
        import matplotlib.pyplot as plt
    except Exception:
        return None

    items = counts.most_common(8)
    labels = [label for label, _count in items]
    values = [count for _label, count in items]
    fig_height = max(2.8, min(6.8, 0.38 * len(items) + 1.8))
    fig, ax = plt.subplots(figsize=(8.0, fig_height), dpi=160)
    fig.patch.set_facecolor("#FFFFFF")
    ax.set_facecolor("#FFFFFF")
    ax.barh(range(len(items)), values, color=color, height=0.55)
    ax.set_title(title, fontsize=12, fontweight="bold", color=THEME["navy"], pad=12)
    ax.set_yticks(range(len(items)))
    ax.set_yticklabels(labels, fontsize=8.2, color="#334155")
    ax.invert_yaxis()
    ax.tick_params(axis="x", labelsize=8, colors="#334155")
    ax.grid(axis="x", color=THEME["border"], linewidth=0.7, alpha=0.65)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.spines["left"].set_color(THEME["border"])
    ax.spines["bottom"].set_color(THEME["border"])
    buffer = io.BytesIO()
    fig.tight_layout()
    fig.savefig(buffer, format="png", bbox_inches="tight")
    plt.close(fig)
    return buffer.getvalue()


def _evidence_distribution_chart_png(rows: list[dict[str, object]]) -> bytes | None:
    return _count_chart_png("Evidence Type Distribution", _evidence_counts(rows), THEME["violet"])


def _indicator_chart_png(rows: list[dict[str, object]]) -> bytes | None:
    return _count_chart_png("Indicator Categories", _indicator_counts(rows), THEME["orange"])


def _status_distribution_chart_png(rows: list[dict[str, object]]) -> bytes | None:
    counts = _risk_counts(rows)
    if not rows:
        return None
    try:
        import matplotlib

        matplotlib.use("Agg", force=True)
        import matplotlib.pyplot as plt
    except Exception:
        return None

    labels = ACTION_STATUSES
    values = [counts.get(label, 0) for label in labels]
    colors = [_prediction_color(label) for label in labels]
    fig, ax = plt.subplots(figsize=(8.0, 2.8), dpi=160)
    fig.patch.set_facecolor("#FFFFFF")
    ax.set_facecolor("#FFFFFF")
    ax.barh(range(len(labels)), values, color=colors, height=0.52)
    ax.set_title("Attention Overview", fontsize=12, fontweight="bold", color=THEME["navy"], pad=10)
    ax.set_yticks(range(len(labels)))
    ax.set_yticklabels(labels, fontsize=8.5, color="#334155")
    ax.invert_yaxis()
    ax.tick_params(axis="x", labelsize=8, colors="#334155")
    ax.grid(axis="x", color=THEME["border"], linewidth=0.7, alpha=0.65)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.spines["left"].set_color(THEME["border"])
    ax.spines["bottom"].set_color(THEME["border"])
    for index, value in enumerate(values):
        ax.text(value + 0.04, index, str(value), va="center", fontsize=8, color="#334155")
    ax.set_xlim(0, max(values + [1]) + 0.7)
    buffer = io.BytesIO()
    fig.tight_layout()
    fig.savefig(buffer, format="png", bbox_inches="tight")
    plt.close(fig)
    return buffer.getvalue()


def _report_profile(incident_context: dict[str, object] | None) -> str:
    selected = _context_text(incident_context, "report_profile", TECHNICAL_PROFILE)
    return STUDENT_PROFILE if selected == STUDENT_PROFILE else TECHNICAL_PROFILE


def _short_explanation(value: object, *, limit: int = 280) -> str:
    text = " ".join(_text(value, "").split())
    if not text:
        return ""
    sentences = [
        sentence.strip()
        for sentence in re.split(r"(?<=[.!?])\s+", text)
        if sentence.strip()
    ]
    concise = " ".join(sentences[:2]) if sentences else text
    return concise if len(concise) <= limit else concise[: limit - 3].rstrip() + "..."


def _row_findings(row: dict[str, object]) -> list[str]:
    bundle_findings = _bundle(row).get("findings", [])
    saved = (
        [_text(item, "") for item in bundle_findings]
        if isinstance(bundle_findings, list)
        else []
    )
    output: list[str] = []
    for item in _flags(row.get("flags")) + saved:
        label = _indicator_label(item)
        if label and label not in output:
            output.append(label)
    if not output and _is_unknown_result(row) and _evidence_family(row) == "Phone":
        output.append("Phone ownership unconfirmed")
    return output


def _relevant_excerpt(row: dict[str, object], *, limit: int = 220) -> str:
    preview = " ".join(_text(row.get("preview"), "").split())
    raw = _text(row.get("raw_input"), "")
    if not raw or (raw.lstrip().startswith("{") and _evidence_family(row) == "Phone"):
        raw = preview
    candidates = [
        " ".join(part.split())
        for part in re.split(r"(?<=[.!?])\s+|[\r\n]+", raw)
        if part.strip()
    ]
    findings = [item.casefold() for item in _row_findings(row)]

    def score(sentence: str) -> tuple[int, int]:
        lowered = sentence.casefold()
        matches = sum(1 for finding in findings if finding and finding in lowered)
        return matches, -len(sentence)

    selected = max(candidates, key=score) if candidates else preview
    if not selected:
        return "No short evidence excerpt was retained."
    return selected if len(selected) <= limit else selected[: limit - 3].rstrip() + "..."


def _student_evidence(row: dict[str, object], index: int) -> dict[str, object]:
    explanation = _short_explanation(row.get("explanation"))
    if not explanation:
        explanation = {
            IMMEDIATE_ACTION: "The available evidence contains signals that require immediate attention.",
            REVIEW_REQUIRED: "The evidence is incomplete, conflicting, or suspicious enough to verify.",
            NO_IMMEDIATE_ACTION: "The available evidence does not indicate immediate action, but normal caution still applies.",
        }.get(_risk_bucket(row), "The available evidence should be reviewed in context.")
    remediation = _remediation_items([row])
    next_action = (
        remediation[0]["action"]
        if remediation
        else _evidence_specific_action(_evidence_family(row))
    )
    return {
        "index": index,
        "family": _evidence_family(row),
        "source": _short_source(row),
        "timestamp": _date_text(row.get("scanned_at")),
        "result": _native_prediction(row),
        "attention": _risk_bucket(row),
        "score": _score_text(row),
        "score_label": _text(row.get("score_label"), "Evidence score"),
        "analysis": _engine_text(row),
        "why": explanation,
        "indicators": _row_findings(row)[:3],
        "excerpt": _relevant_excerpt(row),
        "next_action": next_action,
    }


def _top_actions(
    rows: list[dict[str, object]],
    incident_context: dict[str, object] | None,
    *,
    limit: int = 3,
) -> list[dict[str, str]]:
    return _remediation_items(rows, incident_context)[:limit]


def _unique_artifacts(
    rows: list[dict[str, object]],
    *,
    kind: str | None = None,
) -> list[dict[str, object]]:
    output: list[dict[str, object]] = []
    seen: set[str] = set()
    for row in rows:
        artifacts = _bundle(row).get("artifacts", [])
        if not isinstance(artifacts, list):
            continue
        for artifact in artifacts:
            if not isinstance(artifact, dict):
                continue
            artifact_kind = _text(artifact.get("kind"), "").casefold()
            if kind and artifact_kind != kind.casefold():
                continue
            key = _artifact_figure_key(artifact)
            if not key:
                key = hashlib.sha256(
                    json.dumps(artifact, sort_keys=True, ensure_ascii=True, default=str).encode("utf-8")
                ).hexdigest()
            if key in seen:
                continue
            seen.add(key)
            output.append(artifact)
    return output


SHARED_MODEL_ARTIFACT_TITLES = {
    "performance metrics",
    "performance metrics data",
    "confusion matrix heatmap",
    "roc-auc curve",
}


def _artifact_identity(artifact: dict[str, object]) -> str:
    figure_key = _artifact_figure_key(artifact)
    if figure_key:
        return figure_key
    return hashlib.sha256(
        json.dumps(
            artifact,
            sort_keys=True,
            ensure_ascii=True,
            default=str,
        ).encode("utf-8")
    ).hexdigest()


def _is_shared_model_artifact(artifact: dict[str, object]) -> bool:
    title = " ".join(_text(artifact.get("title"), "").casefold().split())
    return title in SHARED_MODEL_ARTIFACT_TITLES


def _student_artifact_groups(
    rows: list[dict[str, object]],
) -> tuple[
    list[dict[str, object]],
    list[dict[str, object]],
]:
    evidence_groups: list[dict[str, object]] = []
    shared_by_key: dict[str, dict[str, object]] = {}

    for index, row in enumerate(rows, 1):
        evidence_label = f"Evidence {index} - {_evidence_family(row)}"
        bundle_artifacts = _bundle(row).get("artifacts", [])
        artifacts = (
            [item for item in bundle_artifacts if isinstance(item, dict)]
            if isinstance(bundle_artifacts, list)
            else []
        )
        evidence_artifacts: list[dict[str, object]] = []
        local_seen: set[str] = set()
        for artifact in artifacts:
            key = _artifact_identity(artifact)
            if _is_shared_model_artifact(artifact):
                shared_entry = shared_by_key.get(key)
                if shared_entry is None:
                    shared_by_key[key] = {
                        "artifact": artifact,
                        "evidence_refs": [evidence_label],
                    }
                elif evidence_label not in shared_entry["evidence_refs"]:
                    shared_entry["evidence_refs"].append(evidence_label)
                continue
            if key in local_seen:
                continue
            local_seen.add(key)
            evidence_artifacts.append(artifact)
        evidence_groups.append(
            {
                "index": index,
                "row": row,
                "label": evidence_label,
                "artifacts": evidence_artifacts,
            }
        )

    return evidence_groups, list(shared_by_key.values())


def _student_artifact_text_lines(
    artifact: dict[str, object],
    artifact_index: int,
    *,
    evidence_refs: list[str] | None = None,
) -> list[str]:
    kind = _text(artifact.get("kind"), "artifact").title()
    title = _text(artifact.get("title"), f"Artifact {artifact_index}")
    lines = [f"{artifact_index}. [{kind}] {title}"]
    if evidence_refs:
        lines.append(f"   Used by: {', '.join(evidence_refs)}")
    description = _text(artifact.get("description"), "")
    if description:
        lines.append(f"   {description}")
    if artifact.get("figure_sha256"):
        lines.append(f"   Figure SHA-256: {artifact['figure_sha256']}")
    if kind == "Text" and artifact.get("body"):
        lines.append(f"   Content: {_value_text(artifact.get('body'), 5000)}")
    artifact_rows = _artifact_rows(artifact)
    columns = _artifact_columns(artifact, artifact_rows)
    for row_data in artifact_rows:
        lines.append(
            "   - "
            + " | ".join(
                f"{column}: {_value_text(row_data.get(column), 300)}"
                for column in columns
            )
        )
    if kind == "Chart" and not artifact.get("figure_json"):
        lines.append(
            "   Availability: no renderable chart definition was saved for this artifact."
        )
    return lines


def _student_preview_lines(
    rows: list[dict[str, object]],
    report_note: str,
    sections: dict[str, bool],
    incident_context: dict[str, object] | None,
) -> list[str]:
    counts = _risk_counts(rows)
    lines = [
        "AI-FDS Student Investigation Brief",
        f"Generated: {formatted_now()}",
        f"Case: {_context_text(incident_context, 'case_identifier', 'Not assigned')}",
        f"Purpose: {_context_text(incident_context, 'purpose_scope', 'Educational scam-evidence review')}",
        "",
        "Investigation Summary",
        f"- Evidence reviewed: {len(rows)}",
        f"- Immediate action: {counts.get(IMMEDIATE_ACTION, 0)}",
        f"- Review needed: {counts.get(REVIEW_REQUIRED, 0)}",
        f"- No immediate action: {counts.get(NO_IMMEDIATE_ACTION, 0)}",
    ]
    if sections.get("recommendations", True):
        lines.extend(["", "What should I do next?"])
        for index, item in enumerate(_top_actions(rows, incident_context), 1):
            lines.append(f"{index}. [{item['priority']}] {item['action']}")
            lines.append(f"   Why: {item['reason']}")

    if sections.get("evidence", True):
        lines.extend(["", "Evidence Review"])
        for index, row in enumerate(rows, 1):
            evidence = _student_evidence(row, index)
            lines.extend(
                [
                    "",
                    f"Evidence {index} - {evidence['family']} - {evidence['timestamp']}",
                    f"Result: {evidence['result']}",
                    f"Attention: {evidence['attention']}",
                    f"Evidence score: {evidence['score']} ({evidence['score_label']})",
                    f"Analysis method: {evidence['analysis']}",
                    f"Why: {evidence['why']}",
                    "Key indicators: " + (", ".join(evidence["indicators"]) or "No explicit indicators saved"),
                    f"Relevant excerpt: {evidence['excerpt']}",
                    f"Next action: {evidence['next_action']}",
                ]
            )

    if sections.get("explanations", True):
        lines.extend(["", "Combined Investigation Findings"])
        findings = _combined_findings(rows)[:5]
        if findings:
            for item in findings:
                lines.append(
                    f"- {item['Finding']} | {item['Evidence involved']} | {item['Severity']}"
                )
        else:
            lines.append("- No combined indicators were saved.")

    if sections.get("risk", True):
        lines.extend(["", "Visual Evidence Summary", "Attention Overview"])
        for label in ACTION_STATUSES:
            lines.append(f"- {label}: {counts.get(label, 0)}")
        lines.extend(["", "Evidence Type Distribution"])
        for label, count in _evidence_counts(rows).most_common():
            lines.append(f"- {label}: {count}")
        lines.extend(["", "Indicator Categories"])
        indicator_counts = _indicator_counts(rows)
        if indicator_counts:
            for label, count in indicator_counts.most_common(8):
                lines.append(f"- {label}: {count}")
        else:
            lines.append("- No indicator categories were saved.")
        lines.extend(["", "Source-Native Concern Scores"])
        for index, row in enumerate(rows, 1):
            lines.append(
                f"- Evidence {index} - {_evidence_family(row)}: "
                f"{_score_text(row)} ({_text(row.get('score_label'), 'Evidence score')})"
            )
        lines.append(
            "- Interpretation note: scores from email, transcript/audio, and phone evidence use different methods and are not directly comparable."
        )

    lines.extend(["", "Reviewer Note", report_note.strip() or DEFAULT_RECOMMENDATION])

    if sections.get("appendix", False):
        evidence_groups, shared_artifacts = _student_artifact_groups(rows)
        lines.extend(
            [
                "",
                "Technical Appendix",
                "Evidence-Specific Supporting Visuals and Data",
            ]
        )
        for group in evidence_groups:
            row = group["row"]
            evidence = _student_evidence(row, int(group["index"]))
            lines.extend(
                [
                    "",
                    str(group["label"]),
                    f"Result: {evidence['result']}",
                    f"Attention: {evidence['attention']}",
                    f"Source: {evidence['source']}",
                    f"Timestamp: {evidence['timestamp']}",
                    "Supporting Visuals and Data",
                ]
            )
            artifacts = group["artifacts"]
            if artifacts:
                for artifact_index, artifact in enumerate(artifacts, 1):
                    lines.extend(
                        _student_artifact_text_lines(
                            artifact,
                            artifact_index,
                        )
                    )
            else:
                lines.append(
                    "- No investigation-specific dashboard artifacts were captured for this evidence record."
                )

        lines.extend(["", "Shared Model Reference Metrics"])
        if shared_artifacts:
            for artifact_index, entry in enumerate(shared_artifacts, 1):
                lines.extend(
                    _student_artifact_text_lines(
                        entry["artifact"],
                        artifact_index,
                        evidence_refs=list(entry["evidence_refs"]),
                    )
                )
        else:
            lines.append("- No shared model evaluation artifacts were captured.")

        lines.extend(["", "Technical Evidence Register"])
        legacy_count = 0
        for index, row in enumerate(rows, 1):
            provenance = _provenance(row)
            digest = _text(provenance.get("sha256"), "Legacy - not captured")
            if not provenance:
                legacy_count += 1
            lines.append(
                f"- {index} | ID {_text(row.get('id'))} | {_evidence_family(row)} | "
                f"{_date_text(row.get('scanned_at'))} | SHA-256: {digest} | {_engine_text(row)}"
            )
        if legacy_count:
            lines.append(
                f"- Legacy limitation: {legacy_count} selected record(s) predate immutable evidence snapshots."
            )
        lines.extend(["", "Scope and References"])
        lines.append(
            "- Educational scam-awareness support only. Results support review and do not prove identity, intent, legal liability, or authenticity."
        )
        lines.extend(f"- {title}: {url}" for title, url in DOCUMENTATION_REFERENCES)

    return lines


def _filename(extension: str) -> str:
    stamp = now_for_app().strftime("%Y%m%d_%H%M%S")
    return f"AIFDS_Report_{stamp}.{extension.lower()}"


def _overview_lines(rows: list[dict[str, object]]) -> list[str]:
    lines = ["Investigation Summary"]
    lines.extend(f"- {label}: {value}" for label, value in _summary_rows(rows))
    lines.append("")
    lines.append("Evidence Outcome Distribution")
    counts = _risk_counts(rows)
    for label in ACTION_STATUSES:
        lines.append(f"- {label}: {counts.get(label, 0)}")
    return lines


def _selected_evidence_lines(rows: list[dict[str, object]]) -> list[str]:
    lines = ["Selected Evidence Overview"]
    for index, row in enumerate(rows, 1):
        lines.append(
            f"{index}. {_evidence_family(row)} | Source: {_short_source(row)} | "
            f"Native verdict: {_text(row.get('native_prediction') or row.get('prediction'))} | "
            f"Action status: {_risk_bucket(row)} | {_text(row.get('score_label'), 'Concern score')}: {_score_text(row)} | "
            f"Engine: {_engine_text(row)}"
        )
    return lines


def _evidence_specific_action(family: str) -> str:
    actions = {
        "Email": "Verify sender domain, links, attachments, and account requests through official channels.",
        "Transcript": "Check for OTP requests, secrecy, urgency, payment pressure, and impersonation cues.",
        "Audio": "Do not rely on voice familiarity alone; compare voice authenticity with transcript behavior.",
        "Phone": "Carrier metadata does not confirm identity; verify the caller through an official number.",
    }
    return actions.get(family, "Preserve the original evidence and verify through an independent trusted source.")


def _individual_evidence_lines(rows: list[dict[str, object]]) -> list[str]:
    lines = ["Individual Evidence Results"]
    for index, row in enumerate(rows, 1):
        family = _evidence_family(row)
        lines.extend(
            [
                "",
                "-" * 72,
                f"Evidence {index} of {len(rows)} - {family}",
                "-" * 72,
                f"Evidence type: {family}",
                f"Source: {_short_source(row)}",
                f"Native verdict: {_native_prediction(row)}",
                f"Action status: {_risk_bucket(row)}",
                f"{_text(row.get('score_label'), 'Concern score')}: {_score_text(row)}",
                f"Engine: {_engine_text(row)}",
            ]
        )
        flags = ", ".join(_flags(row.get("flags")))
        if flags:
            lines.append(f"Detected indicators: {flags}")
        explanation = _text(row.get("explanation"), "")
        if explanation:
            lines.append(f"Explanation: {explanation}")
        preview = _text(row.get("preview"), "")
        if preview:
            lines.append(f"Evidence preview: {preview[:650]}")
        lines.append("Dashboard summary:")
        lines.extend(f"- {label}: {value}" for label, value in _bundle_summary_rows(row))
        lines.append("Evidence provenance:")
        lines.extend(f"- {label}: {value}" for label, value in _provenance_rows(row))
        bundle = _bundle(row)
        xai = bundle.get("xai", {})
        if isinstance(xai, dict) and xai:
            lines.append("Explainable analysis:")
            lines.append(f"- Method: {_text(xai.get('method'))}")
            lines.append(f"- Scope: {_text(xai.get('scope'))}")
            if xai.get("explanation"):
                lines.append(f"- Explanation: {_text(xai.get('explanation'))}")
            for factor in _xai_rows(row):
                lines.append(
                    "- Factor: "
                    + " | ".join(
                        f"{str(key).replace('_', ' ').title()}: {_value_text(value, 500)}"
                        for key, value in factor.items()
                    )
                )
            limitations = xai.get("limitations", [])
            if isinstance(limitations, list):
                lines.extend(f"- XAI limitation: {_text(item)}" for item in limitations)
        lines.extend(_artifact_lines(row))
    return lines


def _combined_findings_lines(rows: list[dict[str, object]]) -> list[str]:
    lines = ["Combined Investigation Findings"]
    findings = _combined_findings(rows)
    if not findings:
        lines.append("- No combined indicators were available from the selected evidence.")
        return lines
    for item in findings:
        lines.append(
            f"- {item['Finding']} | Evidence involved: {item['Evidence involved']} | "
            f"Severity: {item['Severity']}"
        )
    return lines


def _recommendation_lines(
    rows: list[dict[str, object]],
    report_note: str,
    incident_context: dict[str, object] | None = None,
) -> list[str]:
    note = report_note.strip() or DEFAULT_RECOMMENDATION
    lines = ["Remediation and Prevention Plan"]
    items = _remediation_items(rows, incident_context)
    if not items:
        lines.append("- No trigger-based remediation could be generated from the selected evidence.")
    for index, item in enumerate(items, 1):
        lines.extend(
            [
                f"{index}. [{item['priority']}] {item['action']}",
                f"   Trigger: {item['trigger']}",
                f"   Why: {item['reason']}",
            ]
        )
    lines.extend(["", "Reviewer Note", note])
    return lines


def _scope_lines() -> list[str]:
    lines = [
        "Scope and Limitations",
        "- Email: TF-IDF with trained email classifiers.",
        "- Transcript: multi-model text classification after manual input or Whisper transcription.",
        "- Audio: MFCC voice-authenticity analysis, behavioral audio features, and transcript analysis.",
        "- Phone: Veriphone.io carrier metadata, PenipuMY reputation evidence, and transparent rules.",
        "- This report is an educational capstone prototype output, not legal or forensic proof.",
        "",
        "Documentation References",
    ]
    lines.extend(f"- {title}: {url}" for title, url in DOCUMENTATION_REFERENCES)
    return lines


def build_preview(
    rows: list[dict[str, object]],
    report_note: str,
    sections: dict[str, bool],
    incident_context: dict[str, object] | None = None,
) -> str:
    if _report_profile(incident_context) == STUDENT_PROFILE:
        return "\n".join(
            _student_preview_lines(rows, report_note, sections, incident_context)
        )

    lines = [
        "AI-based Spam and Caller Fraud Detection System",
        "AI Analysis Evidence Report",
        f"Generated: {formatted_now()}",
        f"Case identifier: {_context_text(incident_context, 'case_identifier', 'Not assigned')}",
        f"Requester: {_context_text(incident_context, 'requester', 'Not specified')}",
        f"Report author: {_context_text(incident_context, 'report_author', 'Not specified')}",
        f"Purpose and scope: {_context_text(incident_context, 'purpose_scope', 'Educational scam-evidence review')}",
        f"Evidence disposition: {_context_text(incident_context, 'evidence_disposition', 'Original inputs retained in local investigation history; report files are derivative exports.')}",
        f"Records included: {len(rows)}",
        "",
    ]
    if sections.get("summary", True):
        lines.extend(_overview_lines(rows))
        lines.append("")
    if sections.get("evidence", True):
        lines.extend(_selected_evidence_lines(rows))
        lines.append("")
    if sections.get("explanations", True):
        lines.extend(_individual_evidence_lines(rows))
        lines.append("")
        lines.extend(_combined_findings_lines(rows))
        lines.append("")
    if sections.get("risk", True):
        lines.extend(_confidence_chart_lines(rows))
        lines.append("")
        lines.append("Evidence Type Distribution")
        for label, count in _evidence_counts(rows).most_common():
            lines.append(f"- {label}: {count}")
        lines.append("")
        lines.append("Indicator Categories")
        indicator_counts = _indicator_counts(rows)
        if indicator_counts:
            for label, count in indicator_counts.most_common(8):
                lines.append(f"- {label}: {count}")
        else:
            lines.append("- No indicator categories were saved with the selected evidence.")
        lines.append("")
    if sections.get("recommendations", True):
        lines.extend(_recommendation_lines(rows, report_note, incident_context))
        lines.append("")
    if sections.get("appendix", True):
        lines.extend(_scope_lines())
    return "\n".join(lines)


def build_txt(
    rows: list[dict[str, object]],
    report_note: str,
    sections: dict[str, bool],
    incident_context: dict[str, object] | None = None,
) -> tuple[bytes, str]:
    body = build_preview(rows, report_note, sections, incident_context)
    return body.encode("utf-8"), _filename("txt")


def _build_student_pdf(
    rows: list[dict[str, object]],
    report_note: str,
    sections: dict[str, bool],
    incident_context: dict[str, object] | None,
) -> tuple[bytes, str]:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.platypus import (
        Image,
        KeepTogether,
        PageBreakIfNotEmpty,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )

    artifact_image_cache = (
        _render_artifact_pngs(rows)
        if sections.get("appendix", False)
        else {}
    )
    buffer = io.BytesIO()
    case_identifier = _context_text(
        incident_context,
        "case_identifier",
        "Not assigned",
    )
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=1.55 * cm,
        leftMargin=1.55 * cm,
        topMargin=1.45 * cm,
        bottomMargin=1.45 * cm,
        title="AI-FDS Student Investigation Brief",
    )
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="StudentTitle",
            parent=styles["Title"],
            fontSize=22,
            leading=26,
            textColor=colors.HexColor(THEME["navy"]),
            spaceAfter=3,
        )
    )
    styles.add(
        ParagraphStyle(
            name="StudentSubtitle",
            parent=styles["BodyText"],
            fontSize=10,
            leading=13,
            textColor=colors.HexColor(THEME["cyan"]),
            spaceAfter=10,
        )
    )
    styles.add(
        ParagraphStyle(
            name="StudentHeading",
            parent=styles["Heading2"],
            fontSize=14,
            leading=17,
            textColor=colors.HexColor(THEME["navy"]),
            spaceBefore=8,
            spaceAfter=7,
            keepWithNext=True,
        )
    )
    styles.add(
        ParagraphStyle(
            name="StudentEvidenceHeading",
            parent=styles["Heading3"],
            fontSize=12.2,
            leading=14.5,
            textColor=colors.HexColor(THEME["navy"]),
            spaceBefore=9,
            spaceAfter=5,
            keepWithNext=True,
        )
    )
    styles.add(
        ParagraphStyle(
            name="StudentArtifactGroup",
            parent=styles["Heading3"],
            fontSize=10.6,
            leading=13,
            textColor=colors.HexColor(THEME["violet"]),
            spaceBefore=6,
            spaceAfter=4,
            keepWithNext=True,
        )
    )
    styles.add(
        ParagraphStyle(
            name="StudentArtifactTitle",
            parent=styles["Heading4"],
            fontSize=9.5,
            leading=11.5,
            textColor=colors.HexColor(THEME["blue"]),
            spaceBefore=5,
            spaceAfter=3,
            keepWithNext=True,
        )
    )
    styles.add(
        ParagraphStyle(
            name="StudentCardTitle",
            parent=styles["Heading3"],
            fontSize=10.5,
            leading=13,
            textColor=colors.white,
            spaceAfter=0,
        )
    )
    styles.add(
        ParagraphStyle(
            name="StudentBody",
            parent=styles["BodyText"],
            fontSize=8.8,
            leading=11.5,
            textColor=colors.HexColor("#1E293B"),
            spaceAfter=3,
        )
    )
    styles.add(
        ParagraphStyle(
            name="StudentSmall",
            parent=styles["BodyText"],
            fontSize=7.5,
            leading=9.5,
            textColor=colors.HexColor(THEME["muted"]),
        )
    )
    styles.add(
        ParagraphStyle(
            name="StudentTable",
            parent=styles["BodyText"],
            fontSize=6.7,
            leading=8.2,
            textColor=colors.HexColor("#0F172A"),
        )
    )
    styles.add(
        ParagraphStyle(
            name="StudentTableHeader",
            parent=styles["StudentTable"],
            textColor=colors.white,
        )
    )

    def paragraph(value: object, style_name: str = "StudentBody") -> Paragraph:
        return Paragraph(html.escape(_text(value)), styles[style_name])

    def label_value(label: str, value: object) -> Paragraph:
        return Paragraph(
            f"<b>{html.escape(label)}:</b> {html.escape(_text(value))}",
            styles["StudentBody"],
        )

    def status_fill(status: str) -> str:
        return {
            IMMEDIATE_ACTION: "#FEE2E2",
            REVIEW_REQUIRED: "#FEF3C7",
            NO_IMMEDIATE_ACTION: "#DCFCE7",
        }.get(status, "#F1F5F9")

    def base_table_style(*, header: bool = False) -> TableStyle:
        commands = [
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
            ("RIGHTPADDING", (0, 0), (-1, -1), 6),
            ("TOPPADDING", (0, 0), (-1, -1), 5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ("LINEBELOW", (0, -1), (-1, -1), 0.4, colors.HexColor(THEME["border"])),
        ]
        if header:
            commands.extend(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(THEME["navy"])),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ]
            )
        return TableStyle(commands)

    def add_pdf_artifact(
        artifact: dict[str, object],
        artifact_index: int,
        *,
        evidence_refs: list[str] | None = None,
    ) -> None:
        artifact_kind = _text(artifact.get("kind"), "artifact").title()
        artifact_title = _text(
            artifact.get("title"),
            f"Artifact {artifact_index}",
        )
        artifact_flowables: list[object] = [
            Paragraph(
                html.escape(
                    f"{artifact_index}. [{artifact_kind}] {artifact_title}"
                ),
                styles["StudentArtifactTitle"],
            )
        ]
        if evidence_refs:
            artifact_flowables.append(
                paragraph(
                    f"Used by: {', '.join(evidence_refs)}",
                    "StudentSmall",
                )
            )
        description = _text(artifact.get("description"), "")
        if description:
            artifact_flowables.append(paragraph(description, "StudentSmall"))
        if artifact_kind == "Chart":
            image_bytes = _artifact_png(artifact, artifact_image_cache)
            if image_bytes:
                image = Image(io.BytesIO(image_bytes))
                image._restrictSize(15.8 * cm, 6.2 * cm)
                artifact_flowables.append(image)
            elif artifact.get("figure_json"):
                artifact_flowables.append(
                    paragraph(
                        "The exact chart definition is preserved, but the static image renderer was unavailable for this export."
                    )
                )
            else:
                artifact_flowables.append(
                    paragraph(
                        "This chart was not available when the evidence snapshot was captured."
                    )
                )
        elif artifact_kind == "Text" and artifact.get("body"):
            artifact_flowables.append(
                paragraph(_value_text(artifact.get("body"), 5000))
            )
        story.extend([KeepTogether(artifact_flowables), Spacer(1, 4)])

        artifact_rows = _artifact_rows(artifact)
        columns = _artifact_columns(artifact, artifact_rows)
        for column_group in _chunked_columns(columns):
            table_rows = [
                [
                    paragraph(column, "StudentTableHeader")
                    for column in column_group
                ]
            ]
            table_rows.extend(
                [
                    paragraph(
                        _value_text(row_data.get(column), 900),
                        "StudentTable",
                    )
                    for column in column_group
                ]
                for row_data in artifact_rows
            )
            if len(table_rows) > 1:
                width = 16.8 * cm / len(column_group)
                artifact_table = Table(
                    table_rows,
                    repeatRows=1,
                    colWidths=[width] * len(column_group),
                )
                artifact_table.setStyle(base_table_style(header=True))
                story.extend([artifact_table, Spacer(1, 5)])

    story: list[object] = [
        Paragraph("AI-FDS Student Investigation Brief", styles["StudentTitle"]),
        Paragraph(
            "Plain-language findings, priority actions, and traceable technical evidence",
            styles["StudentSubtitle"],
        ),
    ]
    metadata = Table(
        [
            [
                label_value("Case", case_identifier),
                label_value("Generated", formatted_now()),
            ],
            [
                label_value(
                    "Purpose",
                    _context_text(
                        incident_context,
                        "purpose_scope",
                        "Educational scam-evidence review",
                    ),
                ),
                label_value("Evidence reviewed", len(rows)),
            ],
        ],
        colWidths=[8.8 * cm, 8.0 * cm],
    )
    metadata.setStyle(
        TableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F8FAFC")),
                ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor(THEME["border"])),
                ("INNERGRID", (0, 0), (-1, -1), 0.25, colors.HexColor(THEME["border"])),
                ("LEFTPADDING", (0, 0), (-1, -1), 7),
                ("RIGHTPADDING", (0, 0), (-1, -1), 7),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )
    story.extend([metadata, Spacer(1, 8)])

    if sections.get("summary", True):
        counts = _risk_counts(rows)
        story.append(Paragraph("Investigation Summary", styles["StudentHeading"]))
        summary_values = [
            ("Evidence", len(rows), "#DBEAFE", THEME["blue"]),
            (
                "Immediate",
                counts.get(IMMEDIATE_ACTION, 0),
                "#FEE2E2",
                THEME["red"],
            ),
            (
                "Review",
                counts.get(REVIEW_REQUIRED, 0),
                "#FEF3C7",
                THEME["orange"],
            ),
            (
                "No immediate action",
                counts.get(NO_IMMEDIATE_ACTION, 0),
                "#DCFCE7",
                THEME["green"],
            ),
        ]
        summary_cells = [
            Paragraph(
                f"<font color='{color}'><b>{value}</b></font><br/>"
                f"<font size='7'>{html.escape(label)}</font>",
                styles["StudentBody"],
            )
            for label, value, _fill, color in summary_values
        ]
        summary_table = Table([summary_cells], colWidths=[4.2 * cm] * 4)
        summary_style = [
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("ALIGN", (0, 0), (-1, -1), "CENTER"),
            ("BOX", (0, 0), (-1, -1), 0.45, colors.HexColor(THEME["border"])),
            ("INNERGRID", (0, 0), (-1, -1), 0.45, colors.white),
            ("TOPPADDING", (0, 0), (-1, -1), 8),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
        ]
        for index, (_label, _value, fill, _color) in enumerate(summary_values):
            summary_style.append(
                ("BACKGROUND", (index, 0), (index, 0), colors.HexColor(fill))
            )
        summary_table.setStyle(TableStyle(summary_style))
        story.extend([summary_table, Spacer(1, 8)])

    if sections.get("recommendations", True):
        story.append(Paragraph("What should I do next?", styles["StudentHeading"]))
        actions = _top_actions(rows, incident_context)
        if actions:
            action_rows = []
            for index, item in enumerate(actions, 1):
                action_rows.append(
                    [
                        paragraph(str(index)),
                        Paragraph(
                            f"<b>{html.escape(item['action'])}</b><br/>"
                            f"<font color='{THEME['muted']}'>{html.escape(item['reason'])}</font>",
                            styles["StudentBody"],
                        ),
                        paragraph(item["priority"]),
                    ]
                )
            action_table = Table(
                action_rows,
                colWidths=[0.7 * cm, 13.4 * cm, 2.7 * cm],
            )
            action_table.setStyle(
                TableStyle(
                    [
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F8FAFC")),
                        ("LINEBELOW", (0, 0), (-1, -2), 0.35, colors.HexColor(THEME["border"])),
                        ("LEFTPADDING", (0, 0), (-1, -1), 6),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                        ("TOPPADDING", (0, 0), (-1, -1), 6),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                    ]
                )
            )
            story.extend([action_table, Spacer(1, 8)])
        else:
            story.append(
                paragraph(
                    "No trigger-based action was generated. Preserve the evidence and verify it through an independent official channel."
                )
            )

    if sections.get("evidence", True):
        story.append(Paragraph("Evidence Review", styles["StudentHeading"]))
        for index, row in enumerate(rows, 1):
            evidence = _student_evidence(row, index)
            status = str(evidence["attention"])
            header = Table(
                [
                    [
                        Paragraph(
                            html.escape(
                                f"Evidence {index} - {evidence['family']}"
                            ),
                            styles["StudentCardTitle"],
                        ),
                        Paragraph(
                            f"<b>{html.escape(status)}</b>",
                            styles["StudentSmall"],
                        ),
                    ]
                ],
                colWidths=[11.6 * cm, 5.2 * cm],
            )
            header.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (0, 0), colors.HexColor(THEME["navy"])),
                        ("BACKGROUND", (1, 0), (1, 0), colors.HexColor(status_fill(status))),
                        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                        ("ALIGN", (1, 0), (1, 0), "RIGHT"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 7),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 7),
                        ("TOPPADDING", (0, 0), (-1, -1), 6),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                    ]
                )
            )
            indicators = ", ".join(evidence["indicators"]) or "No explicit indicators saved"
            details = Table(
                [
                    [
                        label_value("Result", evidence["result"]),
                        label_value("Evidence score", evidence["score"]),
                        label_value("Time", evidence["timestamp"]),
                    ],
                    [
                        label_value("Source", evidence["source"]),
                        label_value("Analysis method", evidence["analysis"]),
                        label_value("Score meaning", evidence["score_label"]),
                    ],
                    [
                        Paragraph(
                            f"<b>Why:</b> {html.escape(str(evidence['why']))}<br/>"
                            f"<b>Key indicators:</b> {html.escape(indicators)}<br/>"
                            f"<b>Relevant excerpt:</b> {html.escape(str(evidence['excerpt']))}<br/>"
                            f"<b>Next action:</b> {html.escape(str(evidence['next_action']))}",
                            styles["StudentBody"],
                        )
                    ],
                ],
                colWidths=[5.6 * cm, 5.6 * cm, 5.6 * cm],
            )
            details.setStyle(
                TableStyle(
                    [
                        ("SPAN", (0, 2), (-1, 2)),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("BOX", (0, 0), (-1, -1), 0.45, colors.HexColor(THEME["border"])),
                        ("INNERGRID", (0, 0), (-1, 1), 0.25, colors.HexColor(THEME["border"])),
                        ("BACKGROUND", (0, 0), (-1, 1), colors.HexColor("#F8FAFC")),
                        ("LEFTPADDING", (0, 0), (-1, -1), 7),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 7),
                        ("TOPPADDING", (0, 0), (-1, -1), 6),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                    ]
                )
            )
            story.extend([KeepTogether([header, details]), Spacer(1, 8)])

    if sections.get("explanations", True):
        story.append(
            Paragraph(
                "Combined Investigation Findings",
                styles["StudentHeading"],
            )
        )
        findings = _combined_findings(rows)[:5]
        if findings:
            finding_rows = [
                [
                    paragraph("Finding", "StudentTableHeader"),
                    paragraph("Evidence", "StudentTableHeader"),
                    paragraph("Priority", "StudentTableHeader"),
                ]
            ]
            finding_rows.extend(
                [
                    paragraph(item["Finding"], "StudentTable"),
                    paragraph(item["Evidence involved"], "StudentTable"),
                    paragraph(item["Severity"], "StudentTable"),
                ]
                for item in findings
            )
            finding_table = Table(
                finding_rows,
                repeatRows=1,
                colWidths=[8.0 * cm, 5.5 * cm, 3.3 * cm],
            )
            finding_table.setStyle(base_table_style(header=True))
            story.extend([finding_table, Spacer(1, 7)])
        else:
            story.append(paragraph("No combined indicators were saved."))

    if sections.get("risk", True):
        story.append(
            Paragraph("Visual Evidence Summary", styles["StudentHeading"])
        )
        summary_charts = [
            (_status_distribution_chart_png(rows), 16.0 * cm, 5.5 * cm),
            (_confidence_chart_png(rows), 16.0 * cm, 5.8 * cm),
            (_evidence_distribution_chart_png(rows), 14.8 * cm, 5.0 * cm),
            (_indicator_chart_png(rows), 14.8 * cm, 5.0 * cm),
        ]
        for chart_bytes, max_width, max_height in summary_charts:
            if not chart_bytes:
                continue
            image = Image(io.BytesIO(chart_bytes))
            image._restrictSize(max_width, max_height)
            story.extend([image, Spacer(1, 6)])
        story.append(
            paragraph(
                "Interpretation note: email probabilities, transcript/audio risk signals, and phone evidence scores use different methods and are not directly comparable.",
                "StudentSmall",
            )
        )

    story.append(Paragraph("Reviewer Note", styles["StudentHeading"]))
    story.append(paragraph(report_note.strip() or DEFAULT_RECOMMENDATION))

    if sections.get("appendix", False):
        evidence_groups, shared_artifacts = _student_artifact_groups(rows)
        story.extend(
            [
                PageBreakIfNotEmpty(),
                Paragraph("Technical Appendix", styles["StudentHeading"]),
                paragraph(
                    "Investigation-specific output is grouped under its evidence record. Shared training evaluation metrics appear once in a separate reference section."
                ),
                Paragraph(
                    "Evidence-Specific Supporting Visuals and Data",
                    styles["StudentHeading"],
                ),
            ]
        )

        for group in evidence_groups:
            row = group["row"]
            evidence = _student_evidence(row, int(group["index"]))
            status = str(evidence["attention"])
            evidence_header = Table(
                [
                    [
                        Paragraph(
                            html.escape(str(group["label"])),
                            styles["StudentCardTitle"],
                        ),
                        Paragraph(
                            f"<b>{html.escape(status)}</b>",
                            styles["StudentSmall"],
                        ),
                    ]
                ],
                colWidths=[11.6 * cm, 5.2 * cm],
            )
            evidence_header.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (0, 0), colors.HexColor(THEME["navy"])),
                        ("BACKGROUND", (1, 0), (1, 0), colors.HexColor(status_fill(status))),
                        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                        ("ALIGN", (1, 0), (1, 0), "RIGHT"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 7),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 7),
                        ("TOPPADDING", (0, 0), (-1, -1), 5),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                    ]
                )
            )
            evidence_meta = Table(
                [
                    [
                        label_value("Result", evidence["result"]),
                        label_value("Source", evidence["source"]),
                        label_value("Time", evidence["timestamp"]),
                    ]
                ],
                colWidths=[5.6 * cm, 5.6 * cm, 5.6 * cm],
            )
            evidence_meta.setStyle(
                TableStyle(
                    [
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("BOX", (0, 0), (-1, -1), 0.4, colors.HexColor(THEME["border"])),
                        ("INNERGRID", (0, 0), (-1, -1), 0.25, colors.HexColor(THEME["border"])),
                        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F8FAFC")),
                        ("LEFTPADDING", (0, 0), (-1, -1), 6),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                        ("TOPPADDING", (0, 0), (-1, -1), 5),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                    ]
                )
            )
            story.extend(
                [
                    Spacer(1, 7),
                    KeepTogether(
                        [
                            evidence_header,
                            evidence_meta,
                            Paragraph(
                                "Supporting Visuals and Data",
                                styles["StudentArtifactGroup"],
                            ),
                        ]
                    ),
                ]
            )
            artifacts = group["artifacts"]
            if artifacts:
                for artifact_index, artifact in enumerate(artifacts, 1):
                    add_pdf_artifact(artifact, artifact_index)
            else:
                story.append(
                    paragraph(
                        "No investigation-specific dashboard artifacts were captured for this evidence record."
                    )
                )

        story.append(
            Paragraph(
                "Shared Model Reference Metrics",
                styles["StudentHeading"],
            )
        )
        story.append(
            paragraph(
                "These are model evaluation references, not measurements produced by a single investigation.",
                "StudentSmall",
            )
        )
        if shared_artifacts:
            for artifact_index, entry in enumerate(shared_artifacts, 1):
                add_pdf_artifact(
                    entry["artifact"],
                    artifact_index,
                    evidence_refs=list(entry["evidence_refs"]),
                )
        else:
            story.append(
                paragraph("No shared model evaluation artifacts were captured.")
            )

        story.append(Paragraph("Evidence Register", styles["StudentHeading"]))
        register_rows = [
            [
                paragraph("#", "StudentTableHeader"),
                paragraph("Type", "StudentTableHeader"),
                paragraph("Time", "StudentTableHeader"),
                paragraph("Source", "StudentTableHeader"),
                paragraph("Input SHA-256", "StudentTableHeader"),
                paragraph("AI analysis", "StudentTableHeader"),
            ]
        ]
        legacy_count = 0
        for index, row in enumerate(rows, 1):
            provenance = _provenance(row)
            if not provenance:
                legacy_count += 1
            register_rows.append(
                [
                    paragraph(index, "StudentTable"),
                    paragraph(_evidence_family(row), "StudentTable"),
                    paragraph(_date_text(row.get("scanned_at")), "StudentTable"),
                    paragraph(_short_source(row), "StudentTable"),
                    paragraph(
                        _text(provenance.get("sha256"), "Legacy - not captured"),
                        "StudentTable",
                    ),
                    paragraph(_engine_text(row), "StudentTable"),
                ]
            )
        register = Table(
            register_rows,
            repeatRows=1,
            colWidths=[
                0.6 * cm,
                1.5 * cm,
                2.5 * cm,
                3.0 * cm,
                5.7 * cm,
                3.5 * cm,
            ],
        )
        register.setStyle(base_table_style(header=True))
        story.extend([register, Spacer(1, 6)])
        if legacy_count:
            story.append(
                paragraph(
                    f"Legacy limitation: {legacy_count} selected record(s) predate immutable evidence snapshots."
                )
            )

        story.extend(
            [
                Paragraph("Scope and References", styles["StudentHeading"]),
                paragraph(
                    "Educational scam-awareness support only. Results support review and do not prove identity, intent, legal liability, or authenticity."
                ),
            ]
        )
        for title, url in DOCUMENTATION_REFERENCES:
            story.append(
                Paragraph(
                    f"<b>{html.escape(title)}</b><br/><font size='7'>{html.escape(url)}</font>",
                    styles["StudentSmall"],
                )
            )
            story.append(Spacer(1, 3))

    def page_footer(canvas, document) -> None:
        canvas.saveState()
        canvas.setStrokeColor(colors.HexColor(THEME["border"]))
        canvas.line(1.55 * cm, 1.0 * cm, A4[0] - 1.55 * cm, 1.0 * cm)
        canvas.setFont("Helvetica", 7.2)
        canvas.setFillColor(colors.HexColor(THEME["muted"]))
        canvas.drawString(
            1.55 * cm,
            0.68 * cm,
            f"AI-FDS Student Brief | Case {case_identifier}",
        )
        canvas.drawRightString(
            A4[0] - 1.55 * cm,
            0.68 * cm,
            f"Page {canvas.getPageNumber()}",
        )
        canvas.restoreState()

    doc.build(story, onFirstPage=page_footer, onLaterPages=page_footer)
    return buffer.getvalue(), _filename("pdf")


def build_pdf(
    rows: list[dict[str, object]],
    report_note: str,
    sections: dict[str, bool],
    incident_context: dict[str, object] | None = None,
) -> tuple[bytes, str]:
    if _report_profile(incident_context) == STUDENT_PROFILE:
        return _build_student_pdf(rows, report_note, sections, incident_context)

    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.platypus import (
        Image,
        PageBreakIfNotEmpty,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )

    artifact_image_cache = (
        _render_artifact_pngs(rows)
        if sections.get("explanations", True)
        else {}
    )
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=1.6 * cm,
        leftMargin=1.6 * cm,
        topMargin=1.5 * cm,
        bottomMargin=1.5 * cm,
        title="AI-FDS Analysis Evidence Report",
    )
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="Muted",
            parent=styles["BodyText"],
            fontSize=9,
            leading=12,
            textColor=colors.HexColor("#475569"),
        )
    )
    styles.add(
        ParagraphStyle(
            name="SmallMono",
            parent=styles["BodyText"],
            fontName="Courier",
            fontSize=8,
            leading=10,
            textColor=colors.HexColor("#334155"),
        )
    )
    styles.add(
        ParagraphStyle(
            name="AIFDSHeading",
            parent=styles["Heading2"],
            fontSize=14,
            leading=17,
            textColor=colors.HexColor(THEME["blue"]),
            spaceBefore=10,
            spaceAfter=8,
        )
    )
    styles.add(
        ParagraphStyle(
            name="TinyTable",
            parent=styles["BodyText"],
            fontSize=6.4,
            leading=7.6,
            textColor=colors.HexColor("#0F172A"),
        )
    )
    styles.add(
        ParagraphStyle(
            name="AIFDSSubHeading",
            parent=styles["Heading3"],
            fontSize=10.5,
            leading=13,
            textColor=colors.HexColor(THEME["violet"]),
            spaceBefore=8,
            spaceAfter=5,
        )
    )

    def p(value: object, style_name: str = "Muted") -> Paragraph:
        return Paragraph(html.escape(_text(value)), styles[style_name])

    def table_style(header: bool = True, left_band: str = "soft_blue") -> TableStyle:
        commands = [
            ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor(THEME["border"])),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ("TOPPADDING", (0, 0), (-1, -1), 5),
        ]
        if header:
            commands.extend(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(THEME["navy"])),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ]
            )
        else:
            commands.extend(
                [
                    ("BACKGROUND", (0, 0), (0, -1), colors.HexColor(THEME[left_band])),
                    ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
                ]
            )
        return TableStyle(commands)

    def add_two_column(items: list[tuple[str, str]], left_band: str = "soft_blue") -> None:
        data = [[p(label), p(value)] for label, value in items]
        if not data:
            return
        item_table = Table(data, colWidths=[4.2 * cm, 10.8 * cm])
        item_table.setStyle(table_style(header=False, left_band=left_band))
        story.extend([item_table, Spacer(1, 6)])

    def add_artifact(artifact: dict[str, object], artifact_index: int) -> None:
        title = _text(artifact.get("title"), f"Artifact {artifact_index}")
        kind = _text(artifact.get("kind"), "artifact").title()
        story.append(Paragraph(html.escape(f"{artifact_index}. [{kind}] {title}"), styles["AIFDSSubHeading"]))
        description = _text(artifact.get("description"), "")
        if description:
            story.append(Paragraph(html.escape(description), styles["Muted"]))
        if kind == "Chart":
            image_bytes = _artifact_png(artifact, artifact_image_cache)
            if image_bytes:
                image = Image(io.BytesIO(image_bytes))
                image._restrictSize(15.2 * cm, 9.0 * cm)
                story.extend([image, Spacer(1, 5)])
            elif artifact.get("figure_json"):
                story.append(
                    Paragraph(
                        "The exact Plotly snapshot is preserved in the evidence bundle, but a static image renderer was unavailable for this export.",
                        styles["Muted"],
                    )
                )
            else:
                story.append(Paragraph("Chart unavailable when this investigation was captured.", styles["Muted"]))
        if kind == "Text":
            body = _text(artifact.get("body"), "")
            if body:
                for part in [body[index : index + 3000] for index in range(0, len(body), 3000)]:
                    story.append(
                        Paragraph(
                            html.escape(part).replace("\n", "<br/>"),
                            styles["SmallMono"],
                        )
                    )
        artifact_rows = _artifact_rows(artifact)
        columns = _artifact_columns(artifact, artifact_rows)
        for group_index, column_group in enumerate(_chunked_columns(columns), 1):
            if len(columns) > len(column_group):
                story.append(
                    Paragraph(
                        html.escape(f"Table column group {group_index} of {len(_chunked_columns(columns))}"),
                        styles["Muted"],
                    )
                )
            table_rows = [[p(column) for column in column_group]]
            for row_data in artifact_rows:
                table_rows.append(
                    [p(_value_text(row_data.get(column), 900)) for column in column_group]
                )
            if len(table_rows) > 1:
                width = 15.0 * cm / max(1, len(column_group))
                artifact_table = Table(
                    table_rows,
                    repeatRows=1,
                    colWidths=[width] * len(column_group),
                )
                artifact_table.setStyle(table_style())
                story.extend([artifact_table, Spacer(1, 6)])

    case_identifier = _context_text(incident_context, "case_identifier", "Not assigned")

    story = [
        Paragraph("AI-based Spam and Caller Fraud Detection System", styles["Title"]),
        Paragraph("AI Analysis Evidence Report", styles["AIFDSHeading"]),
        Paragraph(html.escape(f"Generated: {formatted_now()}"), styles["Muted"]),
        Paragraph("Educational capstone prototype. Not legal or forensic proof.", styles["Muted"]),
        Spacer(1, 12),
    ]
    metadata_table = Table(
        [
            [p("Case identifier"), p(case_identifier)],
            [p("Requester"), p(_context_text(incident_context, "requester", "Not specified"))],
            [p("Report author"), p(_context_text(incident_context, "report_author", "Not specified"))],
            [
                p("Purpose and scope"),
                p(_context_text(incident_context, "purpose_scope", "Educational scam-evidence review")),
            ],
            [
                p("Evidence disposition"),
                p(
                    _context_text(
                        incident_context,
                        "evidence_disposition",
                        "Original inputs retained in local investigation history; report files are derivative exports.",
                    )
                ),
            ],
        ],
        colWidths=[4.2 * cm, 10.8 * cm],
    )
    metadata_table.setStyle(table_style(header=False, left_band="soft_green"))
    story.extend([metadata_table, Spacer(1, 12)])

    if sections.get("summary", True):
        story.append(Paragraph("1. Investigation Summary", styles["AIFDSHeading"]))
        summary_rows = [[label, value] for label, value in _summary_rows(rows)]
        table = Table(summary_rows, colWidths=[6 * cm, 9 * cm])
        table.setStyle(table_style(header=False, left_band="soft_violet"))
        story.extend([table, Spacer(1, 12)])
        outcome_chart = _count_chart_png("Evidence Outcome Distribution", _risk_counts(rows), THEME["blue"])
        if outcome_chart:
            story.append(Image(io.BytesIO(outcome_chart), width=14.2 * cm, height=5.2 * cm))
            story.append(Spacer(1, 10))

    if sections.get("evidence", True):
        story.append(Paragraph("2. Selected Evidence Overview", styles["AIFDSHeading"]))
        table_rows = [["#", "Evidence", "Source", "Native verdict", "Action status", "Concern", "Engine"]]
        for index, row in enumerate(rows, 1):
            table_rows.append(
                [
                    Paragraph(str(index), styles["TinyTable"]),
                    Paragraph(html.escape(_evidence_family(row)), styles["TinyTable"]),
                    Paragraph(html.escape(_short_source(row)), styles["TinyTable"]),
                    Paragraph(html.escape(_native_prediction(row)), styles["TinyTable"]),
                    Paragraph(html.escape(_risk_bucket(row)), styles["TinyTable"]),
                    Paragraph(html.escape(_score_text(row)), styles["TinyTable"]),
                    Paragraph(html.escape(_engine_text(row)), styles["TinyTable"]),
                ]
            )
        table = Table(
            table_rows,
            repeatRows=1,
            colWidths=[0.6 * cm, 1.4 * cm, 2.3 * cm, 2.1 * cm, 3.2 * cm, 1.5 * cm, 3.9 * cm],
        )
        table.setStyle(table_style())
        story.extend([table, Spacer(1, 12)])

    if sections.get("explanations", True):
        story.append(Paragraph("3. Individual Evidence Results", styles["AIFDSHeading"]))
        for index, row in enumerate(rows, 1):
            family = _evidence_family(row)
            story.append(Paragraph(html.escape(f"Evidence {index} of {len(rows)} - {family}"), styles["AIFDSSubHeading"]))
            profile_rows = [
                ["Evidence type", family],
                ["Source", _short_source(row)],
                ["Native verdict", _native_prediction(row)],
                ["Action status", _risk_bucket(row)],
                [_text(row.get("score_label"), "Concern score"), _score_text(row)],
                ["Engine", _engine_text(row)],
                ["Timestamp", _date_text(row.get("scanned_at"))],
            ]
            profile_table = Table(profile_rows, colWidths=[4 * cm, 11 * cm])
            profile_table.setStyle(table_style(header=False, left_band="soft_blue"))
            story.extend([profile_table, Spacer(1, 5)])
            flags = ", ".join(_flags(row.get("flags")))
            if flags:
                story.append(Paragraph(html.escape(f"Detected indicators: {flags}"), styles["Muted"]))
            explanation = _text(row.get("explanation"), "")
            if explanation:
                story.append(Paragraph(html.escape(explanation), styles["Muted"]))
            preview = _text(row.get("preview"), "")
            if preview:
                story.append(Paragraph(html.escape(preview[:350]), styles["SmallMono"]))
            story.append(Paragraph("Dashboard Summary", styles["AIFDSSubHeading"]))
            add_two_column(_bundle_summary_rows(row), left_band="soft_violet")
            story.append(Paragraph("Evidence Provenance", styles["AIFDSSubHeading"]))
            add_two_column(_provenance_rows(row), left_band="soft_green")
            bundle = _bundle(row)
            xai = bundle.get("xai", {})
            if isinstance(xai, dict) and xai:
                story.append(Paragraph("Explainable Analysis", styles["AIFDSSubHeading"]))
                add_two_column(
                    [
                        ("Method", _text(xai.get("method"))),
                        ("Scope", _text(xai.get("scope"))),
                        ("Explanation", _text(xai.get("explanation"))),
                    ],
                    left_band="soft_orange",
                )
                factors = _xai_rows(row)
                if factors:
                    factor_columns = _artifact_columns({}, factors)
                    for column_group in _chunked_columns(factor_columns):
                        factor_rows = [[p(column) for column in column_group]]
                        for factor in factors:
                            factor_rows.append(
                                [p(_value_text(factor.get(column), 600)) for column in column_group]
                            )
                        width = 15.0 * cm / max(1, len(column_group))
                        factor_table = Table(
                            factor_rows,
                            repeatRows=1,
                            colWidths=[width] * len(column_group),
                        )
                        factor_table.setStyle(table_style())
                        story.extend([factor_table, Spacer(1, 5)])
                for limitation in xai.get("limitations", []) if isinstance(xai.get("limitations"), list) else []:
                    story.append(Paragraph(html.escape(f"XAI limitation: {_text(limitation)}"), styles["Muted"]))
            artifacts = bundle.get("artifacts", [])
            if isinstance(artifacts, list) and artifacts:
                story.append(Paragraph("Saved Dashboard Artifacts", styles["AIFDSSubHeading"]))
                for artifact_index, artifact in enumerate(artifacts, 1):
                    if isinstance(artifact, dict):
                        add_artifact(artifact, artifact_index)
            else:
                story.append(
                    Paragraph(
                        "Legacy record: no ordered dashboard artifact snapshot was saved.",
                        styles["Muted"],
                    )
                )
            story.append(Spacer(1, 8))

        combined = _combined_findings(rows)
        story.append(Paragraph("Combined Investigation Findings", styles["AIFDSSubHeading"]))
        if combined:
            table_rows = [["Finding", "Evidence involved", "Severity"]] + [
                [item["Finding"], item["Evidence involved"], item["Severity"]] for item in combined
            ]
            table = Table(table_rows, repeatRows=1, colWidths=[6 * cm, 6 * cm, 3 * cm])
            table.setStyle(table_style())
            story.extend([table, Spacer(1, 12)])
        else:
            story.append(Paragraph("No combined indicators were available from the selected evidence.", styles["Muted"]))

    if sections.get("risk", True):
        story.append(Paragraph("4. Visual Evidence Summary", styles["AIFDSHeading"]))
        chart = _confidence_chart_png(rows)
        if chart:
            story.append(Image(io.BytesIO(chart), width=15.2 * cm, height=6.2 * cm))
            story.append(Spacer(1, 8))
        distribution_chart = _evidence_distribution_chart_png(rows)
        if distribution_chart:
            story.append(Image(io.BytesIO(distribution_chart), width=14.2 * cm, height=5.2 * cm))
            story.append(Spacer(1, 8))
        indicator_chart = _indicator_chart_png(rows)
        if indicator_chart:
            story.append(Image(io.BytesIO(indicator_chart), width=14.2 * cm, height=5.2 * cm))
            story.append(Spacer(1, 8))
        story.append(
            Paragraph(
                "The horizontal chart uses each source's explicitly labelled concern score where available. "
                "Model probability, audio reliability-weighted evidence, and phone rule scores remain distinct native measures.",
                styles["Muted"],
            )
        )
        story.append(Spacer(1, 12))

    if sections.get("recommendations", True):
        story.append(Paragraph("5. Remediation and Prevention Plan", styles["AIFDSHeading"]))
        for item in _remediation_items(rows, incident_context):
            story.append(
                Paragraph(
                    html.escape(f"[{item['priority']}] {item['action']}"),
                    styles["AIFDSSubHeading"],
                )
            )
            story.append(Paragraph(html.escape(f"Trigger: {item['trigger']}"), styles["Muted"]))
            story.append(Paragraph(html.escape(f"Why: {item['reason']}"), styles["Muted"]))
        story.append(Paragraph("Reviewer Note", styles["AIFDSSubHeading"]))
        story.append(Paragraph(html.escape(report_note.strip() or DEFAULT_RECOMMENDATION), styles["Muted"]))
        story.append(Spacer(1, 12))

    if sections.get("appendix", True):
        story.append(PageBreakIfNotEmpty())
        story.append(Paragraph("6. Scope and Limitations", styles["AIFDSHeading"]))
        appendix = [
            ["System", "AI-FDS Capstone Prototype"],
            ["Email", "TF-IDF with trained email classifiers"],
            ["Transcript", "Multi-model text classification after manual input or Whisper transcription"],
            ["Audio", "MFCC voice-authenticity analysis, behavioral audio features, and transcript analysis"],
            ["Phone", "Veriphone.io carrier metadata, PenipuMY reputation evidence, and transparent rules"],
            ["Scope", "Educational scam awareness support. Not enterprise security, telecom verification, or legal evidence."],
        ]
        table = Table(appendix, colWidths=[4 * cm, 11 * cm])
        table.setStyle(table_style(header=False, left_band="soft_green"))
        story.extend([table, Spacer(1, 12)])
        story.append(Paragraph("Documentation References", styles["AIFDSSubHeading"]))
        for title, url in DOCUMENTATION_REFERENCES:
            story.append(Paragraph(html.escape(f"{title}: {url}"), styles["Muted"]))

    def page_footer(canvas, document) -> None:
        canvas.saveState()
        canvas.setStrokeColor(colors.HexColor(THEME["border"]))
        canvas.line(1.6 * cm, 1.05 * cm, A4[0] - 1.6 * cm, 1.05 * cm)
        canvas.setFont("Helvetica", 7.5)
        canvas.setFillColor(colors.HexColor(THEME["muted"]))
        canvas.drawString(1.6 * cm, 0.72 * cm, f"AI-FDS Evidence Report | Case {case_identifier}")
        canvas.drawRightString(A4[0] - 1.6 * cm, 0.72 * cm, f"Page {canvas.getPageNumber()}")
        canvas.restoreState()

    doc.build(story, onFirstPage=page_footer, onLaterPages=page_footer)
    return buffer.getvalue(), _filename("pdf")


def _build_student_docx(
    rows: list[dict[str, object]],
    report_note: str,
    sections: dict[str, bool],
    incident_context: dict[str, object] | None,
) -> tuple[bytes, str]:
    from docx import Document
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    artifact_image_cache = (
        _render_artifact_pngs(rows)
        if sections.get("appendix", False)
        else {}
    )
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.58)
    section.bottom_margin = Inches(0.58)
    section.left_margin = Inches(0.62)
    section.right_margin = Inches(0.62)
    case_identifier = _context_text(
        incident_context,
        "case_identifier",
        "Not assigned",
    )

    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9)
    for style_name, size, color in [
        ("Title", 22, "0B1220"),
        ("Heading 1", 15, "0B1220"),
        ("Heading 2", 11, "2563EB"),
        ("Heading 3", 9.5, "7C3AED"),
    ]:
        style = document.styles[style_name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)

    def shade_cell(
        cell,
        fill: str,
        *,
        color: str | None = None,
        bold: bool = False,
    ) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), fill)
        tc_pr.append(shading)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for cell_paragraph in cell.paragraphs:
            for run in cell_paragraph.runs:
                if color:
                    run.font.color.rgb = RGBColor.from_string(color)
                run.font.bold = bold

    def set_repeat_header(row) -> None:
        tr_pr = row._tr.get_or_add_trPr()
        repeat = OxmlElement("w:tblHeader")
        repeat.set(qn("w:val"), "true")
        tr_pr.append(repeat)

    def prevent_row_split(row) -> None:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)

    def set_cell_text(
        cell,
        text: object,
        *,
        bold: bool = False,
        color: str | None = None,
        size: float = 8.5,
    ) -> None:
        cell.text = _text(text)
        for cell_paragraph in cell.paragraphs:
            cell_paragraph.paragraph_format.space_after = Pt(0)
            for run in cell_paragraph.runs:
                run.font.name = "Aptos"
                run.font.size = Pt(size)
                run.font.bold = bold
                if color:
                    run.font.color.rgb = RGBColor.from_string(color)

    def add_heading(text: str, level: int = 1) -> None:
        heading = document.add_heading(text, level=level)
        heading.paragraph_format.space_before = Pt(7)
        heading.paragraph_format.space_after = Pt(4)

    def add_body(text: object, *, muted: bool = False) -> None:
        body = document.add_paragraph(_text(text))
        body.paragraph_format.space_after = Pt(4)
        for run in body.runs:
            run.font.name = "Aptos"
            run.font.size = Pt(8.8)
            if muted:
                run.font.color.rgb = RGBColor.from_string("475569")

    def add_labelled_paragraph(label: str, value: object) -> None:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(3)
        label_run = paragraph.add_run(f"{label}: ")
        label_run.bold = True
        paragraph.add_run(_text(value))
        for run in paragraph.runs:
            run.font.name = "Aptos"
            run.font.size = Pt(8.8)

    def add_picture_centered(image_bytes: bytes, *, width: float) -> None:
        document.add_picture(io.BytesIO(image_bytes), width=Inches(width))
        document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_docx_artifact(
        artifact: dict[str, object],
        artifact_index: int,
        *,
        evidence_refs: list[str] | None = None,
    ) -> None:
        artifact_kind = _text(artifact.get("kind"), "artifact").title()
        artifact_title = _text(
            artifact.get("title"),
            f"Artifact {artifact_index}",
        )
        add_heading(
            f"{artifact_index}. [{artifact_kind}] {artifact_title}",
            level=3,
        )
        if evidence_refs:
            add_body(f"Used by: {', '.join(evidence_refs)}", muted=True)
        description = _text(artifact.get("description"), "")
        if description:
            add_body(description, muted=True)
        if artifact_kind == "Chart":
            image_bytes = _artifact_png(artifact, artifact_image_cache)
            if image_bytes:
                add_picture_centered(image_bytes, width=6.15)
            elif artifact.get("figure_json"):
                add_body(
                    "The exact chart definition is preserved, but the static image renderer was unavailable for this export."
                )
            else:
                add_body(
                    "This chart was not available when the evidence snapshot was captured."
                )
        elif artifact_kind == "Text" and artifact.get("body"):
            add_body(_value_text(artifact.get("body"), 5000))

        artifact_rows = _artifact_rows(artifact)
        columns = _artifact_columns(artifact, artifact_rows)
        for column_group in _chunked_columns(columns):
            artifact_table = document.add_table(
                rows=1,
                cols=len(column_group),
            )
            artifact_table.style = "Table Grid"
            for cell, column in zip(
                artifact_table.rows[0].cells,
                column_group,
            ):
                set_cell_text(cell, column, bold=True, color="FFFFFF", size=7)
                shade_cell(cell, "0B1220", color="FFFFFF", bold=True)
            set_repeat_header(artifact_table.rows[0])
            for row_data in artifact_rows:
                cells = artifact_table.add_row().cells
                for cell, column in zip(cells, column_group):
                    set_cell_text(
                        cell,
                        _value_text(row_data.get(column), 900),
                        size=7,
                    )

    title = document.add_heading("AI-FDS Student Investigation Brief", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(2)
    subtitle = document.add_paragraph(
        "Plain-language findings, priority actions, and traceable technical evidence"
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(8)
    if subtitle.runs:
        subtitle.runs[0].font.color.rgb = RGBColor.from_string("0891B2")
        subtitle.runs[0].font.size = Pt(9.5)

    metadata = document.add_table(rows=2, cols=2)
    metadata.style = "Table Grid"
    metadata.alignment = WD_TABLE_ALIGNMENT.CENTER
    metadata_values = [
        ("Case", case_identifier),
        ("Generated", formatted_now()),
        (
            "Purpose",
            _context_text(
                incident_context,
                "purpose_scope",
                "Educational scam-evidence review",
            ),
        ),
        ("Evidence reviewed", len(rows)),
    ]
    for cell, (label, value) in zip(
        [cell for table_row in metadata.rows for cell in table_row.cells],
        metadata_values,
    ):
        cell.text = ""
        paragraph = cell.paragraphs[0]
        label_run = paragraph.add_run(f"{label}: ")
        label_run.bold = True
        paragraph.add_run(_text(value))
        shade_cell(cell, "F8FAFC")
        for run in paragraph.runs:
            run.font.name = "Aptos"
            run.font.size = Pt(8.5)

    if sections.get("summary", True):
        counts = _risk_counts(rows)
        add_heading("Investigation Summary")
        summary = document.add_table(rows=1, cols=4)
        summary.style = "Table Grid"
        summary.alignment = WD_TABLE_ALIGNMENT.CENTER
        values = [
            ("Evidence", len(rows), "DBEAFE", "2563EB"),
            (
                "Immediate",
                counts.get(IMMEDIATE_ACTION, 0),
                "FEE2E2",
                "DC2626",
            ),
            (
                "Review",
                counts.get(REVIEW_REQUIRED, 0),
                "FEF3C7",
                "D97706",
            ),
            (
                "No immediate action",
                counts.get(NO_IMMEDIATE_ACTION, 0),
                "DCFCE7",
                "059669",
            ),
        ]
        for cell, (label, value, fill, color) in zip(summary.rows[0].cells, values):
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            value_run = paragraph.add_run(f"{value}\n")
            value_run.bold = True
            value_run.font.size = Pt(15)
            value_run.font.color.rgb = RGBColor.from_string(color)
            label_run = paragraph.add_run(label)
            label_run.font.size = Pt(7.5)
            shade_cell(cell, fill)

    if sections.get("recommendations", True):
        add_heading("What should I do next?")
        actions = _top_actions(rows, incident_context)
        for index, item in enumerate(actions, 1):
            action = document.add_table(rows=1, cols=3)
            action.style = "Light Shading Accent 1"
            set_cell_text(action.cell(0, 0), index, bold=True, color="2563EB")
            action.cell(0, 1).text = ""
            action_paragraph = action.cell(0, 1).paragraphs[0]
            action_run = action_paragraph.add_run(item["action"])
            action_run.bold = True
            action_paragraph.add_run(f"\n{item['reason']}")
            for run in action_paragraph.runs:
                run.font.name = "Aptos"
                run.font.size = Pt(8.5)
            set_cell_text(action.cell(0, 2), item["priority"], bold=True)
            prevent_row_split(action.rows[0])

    if sections.get("evidence", True):
        add_heading("Evidence Review")
        for index, row in enumerate(rows, 1):
            evidence = _student_evidence(row, index)
            status = str(evidence["attention"])
            status_fill = {
                IMMEDIATE_ACTION: "FEE2E2",
                REVIEW_REQUIRED: "FEF3C7",
                NO_IMMEDIATE_ACTION: "DCFCE7",
            }.get(status, "F1F5F9")
            header = document.add_table(rows=1, cols=2)
            header.style = "Table Grid"
            set_cell_text(
                header.cell(0, 0),
                f"Evidence {index} - {evidence['family']}",
                bold=True,
                color="FFFFFF",
                size=10,
            )
            shade_cell(header.cell(0, 0), "0B1220", color="FFFFFF", bold=True)
            set_cell_text(header.cell(0, 1), status, bold=True)
            shade_cell(header.cell(0, 1), status_fill, bold=True)
            prevent_row_split(header.rows[0])

            facts = document.add_table(rows=2, cols=3)
            facts.style = "Table Grid"
            fact_values = [
                ("Result", evidence["result"]),
                ("Evidence score", evidence["score"]),
                ("Time", evidence["timestamp"]),
                ("Source", evidence["source"]),
                ("Analysis method", evidence["analysis"]),
                ("Score meaning", evidence["score_label"]),
            ]
            for cell, (label, value) in zip(
                [cell for table_row in facts.rows for cell in table_row.cells],
                fact_values,
            ):
                cell.text = ""
                paragraph = cell.paragraphs[0]
                label_run = paragraph.add_run(f"{label}: ")
                label_run.bold = True
                paragraph.add_run(_text(value))
                shade_cell(cell, "F8FAFC")
                for run in paragraph.runs:
                    run.font.name = "Aptos"
                    run.font.size = Pt(8)
            indicators = ", ".join(evidence["indicators"]) or "No explicit indicators saved"
            details = document.add_table(rows=1, cols=1)
            details.style = "Table Grid"
            details.cell(0, 0).text = ""
            detail_paragraph = details.cell(0, 0).paragraphs[0]
            detail_values = [
                ("Why", evidence["why"]),
                ("Key indicators", indicators),
                ("Relevant excerpt", evidence["excerpt"]),
                ("Next action", evidence["next_action"]),
            ]
            for detail_index, (label, value) in enumerate(detail_values):
                label_run = detail_paragraph.add_run(f"{label}: ")
                label_run.bold = True
                detail_paragraph.add_run(_text(value))
                if detail_index < len(detail_values) - 1:
                    detail_paragraph.add_run("\n")
            for run in detail_paragraph.runs:
                run.font.name = "Aptos"
                run.font.size = Pt(8.3)
            prevent_row_split(details.rows[0])
            document.add_paragraph().paragraph_format.space_after = Pt(1)

    if sections.get("explanations", True):
        add_heading("Combined Investigation Findings")
        findings = _combined_findings(rows)[:5]
        if findings:
            finding_table = document.add_table(rows=1, cols=3)
            finding_table.style = "Table Grid"
            for cell, label in zip(
                finding_table.rows[0].cells,
                ["Finding", "Evidence", "Priority"],
            ):
                set_cell_text(cell, label, bold=True, color="FFFFFF")
                shade_cell(cell, "0B1220", color="FFFFFF", bold=True)
            set_repeat_header(finding_table.rows[0])
            for item in findings:
                cells = finding_table.add_row().cells
                for cell, value in zip(
                    cells,
                    [
                        item["Finding"],
                        item["Evidence involved"],
                        item["Severity"],
                    ],
                ):
                    set_cell_text(cell, value, size=8)
        else:
            add_body("No combined indicators were saved.")

    if sections.get("risk", True):
        add_heading("Visual Evidence Summary")
        summary_charts = [
            (_status_distribution_chart_png(rows), 6.2),
            (_confidence_chart_png(rows), 6.2),
            (_evidence_distribution_chart_png(rows), 5.9),
            (_indicator_chart_png(rows), 5.9),
        ]
        for chart_bytes, width in summary_charts:
            if chart_bytes:
                add_picture_centered(chart_bytes, width=width)
        add_body(
            "Interpretation note: email probabilities, transcript/audio risk signals, and phone evidence scores use different methods and are not directly comparable.",
            muted=True,
        )

    add_heading("Reviewer Note")
    add_body(report_note.strip() or DEFAULT_RECOMMENDATION)

    if sections.get("appendix", False):
        evidence_groups, shared_artifacts = _student_artifact_groups(rows)
        document.add_page_break()
        add_heading("Technical Appendix")
        add_body(
            "Investigation-specific output is grouped under its evidence record. Shared training evaluation metrics appear once in a separate reference section.",
            muted=True,
        )
        add_heading("Evidence-Specific Supporting Visuals and Data")

        for group in evidence_groups:
            row = group["row"]
            evidence = _student_evidence(row, int(group["index"]))
            status = str(evidence["attention"])
            status_fill = {
                IMMEDIATE_ACTION: "FEE2E2",
                REVIEW_REQUIRED: "FEF3C7",
                NO_IMMEDIATE_ACTION: "DCFCE7",
            }.get(status, "F1F5F9")
            header = document.add_table(rows=1, cols=2)
            header.style = "Table Grid"
            set_cell_text(
                header.cell(0, 0),
                str(group["label"]),
                bold=True,
                color="FFFFFF",
                size=10,
            )
            shade_cell(
                header.cell(0, 0),
                "0B1220",
                color="FFFFFF",
                bold=True,
            )
            set_cell_text(header.cell(0, 1), status, bold=True)
            shade_cell(header.cell(0, 1), status_fill, bold=True)
            prevent_row_split(header.rows[0])

            facts = document.add_table(rows=1, cols=3)
            facts.style = "Table Grid"
            for cell, (label, value) in zip(
                facts.rows[0].cells,
                [
                    ("Result", evidence["result"]),
                    ("Source", evidence["source"]),
                    ("Time", evidence["timestamp"]),
                ],
            ):
                cell.text = ""
                paragraph = cell.paragraphs[0]
                label_run = paragraph.add_run(f"{label}: ")
                label_run.bold = True
                paragraph.add_run(_text(value))
                shade_cell(cell, "F8FAFC")
                for run in paragraph.runs:
                    run.font.name = "Aptos"
                    run.font.size = Pt(8)
            add_heading("Supporting Visuals and Data", level=2)
            artifacts = group["artifacts"]
            if artifacts:
                for artifact_index, artifact in enumerate(artifacts, 1):
                    add_docx_artifact(artifact, artifact_index)
            else:
                add_body(
                    "No investigation-specific dashboard artifacts were captured for this evidence record."
                )

        add_heading("Shared Model Reference Metrics")
        add_body(
            "These are model evaluation references, not measurements produced by a single investigation.",
            muted=True,
        )
        if shared_artifacts:
            for artifact_index, entry in enumerate(shared_artifacts, 1):
                add_docx_artifact(
                    entry["artifact"],
                    artifact_index,
                    evidence_refs=list(entry["evidence_refs"]),
                )
        else:
            add_body("No shared model evaluation artifacts were captured.")

        add_heading("Evidence Register")
        register = document.add_table(rows=1, cols=6)
        register.style = "Table Grid"
        for cell, label in zip(
            register.rows[0].cells,
            ["#", "Type", "Time", "Source", "Input SHA-256", "AI analysis"],
        ):
            set_cell_text(cell, label, bold=True, color="FFFFFF", size=7)
            shade_cell(cell, "0B1220", color="FFFFFF", bold=True)
        set_repeat_header(register.rows[0])
        legacy_count = 0
        for index, row in enumerate(rows, 1):
            provenance = _provenance(row)
            if not provenance:
                legacy_count += 1
            values = [
                index,
                _evidence_family(row),
                _date_text(row.get("scanned_at")),
                _short_source(row),
                _text(provenance.get("sha256"), "Legacy - not captured"),
                _engine_text(row),
            ]
            cells = register.add_row().cells
            for cell, value in zip(cells, values):
                set_cell_text(cell, value, size=6.7)
        if legacy_count:
            add_body(
                f"Legacy limitation: {legacy_count} selected record(s) predate immutable evidence snapshots.",
                muted=True,
            )

        add_heading("Scope and References")
        add_body(
            "Educational scam-awareness support only. Results support review and do not prove identity, intent, legal liability, or authenticity."
        )
        for title, url in DOCUMENTATION_REFERENCES:
            add_labelled_paragraph(title, url)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run(
        f"AI-FDS Student Brief | Case {case_identifier} | Page "
    )
    footer_run.font.size = Pt(7)
    page_field = OxmlElement("w:fldSimple")
    page_field.set(qn("w:instr"), "PAGE")
    footer._p.append(page_field)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue(), _filename("docx")


def build_docx(
    rows: list[dict[str, object]],
    report_note: str,
    sections: dict[str, bool],
    incident_context: dict[str, object] | None = None,
) -> tuple[bytes, str]:
    if _report_profile(incident_context) == STUDENT_PROFILE:
        return _build_student_docx(
            rows,
            report_note,
            sections,
            incident_context,
        )

    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    artifact_image_cache = (
        _render_artifact_pngs(rows)
        if sections.get("explanations", True)
        else {}
    )
    document = Document()
    title = document.add_heading("AI-based Spam and Caller Fraud Detection System", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if title.runs:
        title.runs[0].font.color.rgb = RGBColor(37, 99, 235)
    subtitle = document.add_paragraph("AI Analysis Evidence Report")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if subtitle.runs:
        subtitle.runs[0].font.color.rgb = RGBColor(124, 58, 237)
        subtitle.runs[0].font.bold = True
    case_identifier = _context_text(incident_context, "case_identifier", "Not assigned")
    meta = document.add_paragraph(
        f"Generated: {formatted_now()}\n"
        f"Case identifier: {case_identifier}\n"
        f"Requester: {_context_text(incident_context, 'requester', 'Not specified')}\n"
        f"Report author: {_context_text(incident_context, 'report_author', 'Not specified')}\n"
        f"Purpose and scope: {_context_text(incident_context, 'purpose_scope', 'Educational scam-evidence review')}\n"
        f"Evidence disposition: {_context_text(incident_context, 'evidence_disposition', 'Original inputs retained in local investigation history; report files are derivative exports.')}\n"
        f"Records included: {len(rows)}\n"
        "Educational capstone prototype. Not legal or forensic proof."
    )
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER

    footer = document.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run(f"AI-FDS Evidence Report | Case {case_identifier} | Page ")
    page_field = OxmlElement("w:fldSimple")
    page_field.set(qn("w:instr"), "PAGE")
    footer._p.append(page_field)

    def heading(text: str, level: int = 1) -> None:
        paragraph = document.add_heading(text, level=level)
        if paragraph.runs:
            paragraph.runs[0].font.color.rgb = (
                RGBColor(37, 99, 235)
                if level == 1
                else RGBColor(124, 58, 237)
            )

    def body(text: str) -> None:
        paragraph = document.add_paragraph(text)
        for run in paragraph.runs:
            run.font.size = Pt(10)

    def shade_cell(cell, fill: str, text_color: RGBColor | None = None, bold: bool = False) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), fill)
        tc_pr.append(shading)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                if text_color is not None:
                    run.font.color.rgb = text_color
                run.font.bold = bold

    def set_table_header(table) -> None:
        for cell in table.rows[0].cells:
            shade_cell(cell, "0B1220", RGBColor(255, 255, 255), True)

    def two_column_table(items: list[tuple[str, str]]) -> None:
        table = document.add_table(rows=0, cols=2)
        table.style = "Table Grid"
        for label, value in items:
            cells = table.add_row().cells
            cells[0].text = label
            cells[1].text = value
            shade_cell(cells[0], "EDE9FE", RGBColor(15, 23, 42), True)

    def artifact_table(
        rows_data: list[dict[str, object]],
        columns: list[str],
    ) -> None:
        for group_index, column_group in enumerate(_chunked_columns(columns), 1):
            if len(columns) > len(column_group):
                body(f"Table column group {group_index} of {len(_chunked_columns(columns))}")
            table = document.add_table(rows=1, cols=len(column_group))
            table.style = "Table Grid"
            for column_index, column in enumerate(column_group):
                table.rows[0].cells[column_index].text = column
            set_table_header(table)
            for row_data in rows_data:
                cells = table.add_row().cells
                for column_index, column in enumerate(column_group):
                    cells[column_index].text = _value_text(row_data.get(column), 1200)

    def add_artifact(artifact: dict[str, object], artifact_index: int) -> None:
        artifact_kind = _text(artifact.get("kind"), "artifact").title()
        artifact_title = _text(artifact.get("title"), f"Artifact {artifact_index}")
        heading(f"{artifact_index}. [{artifact_kind}] {artifact_title}", level=2)
        description = _text(artifact.get("description"), "")
        if description:
            body(description)
        if artifact_kind == "Chart":
            image_bytes = _artifact_png(artifact, artifact_image_cache)
            if image_bytes:
                document.add_picture(io.BytesIO(image_bytes), width=Inches(6.35))
            elif artifact.get("figure_json"):
                body(
                    "The exact Plotly snapshot is preserved in the evidence bundle, "
                    "but a static image renderer was unavailable for this export."
                )
            else:
                body("Chart unavailable when this investigation was captured.")
        if artifact_kind == "Text" and artifact.get("body"):
            paragraph = document.add_paragraph(str(artifact.get("body", "")))
            for run in paragraph.runs:
                run.font.name = "Courier New"
                run.font.size = Pt(8)
        rows_data = _artifact_rows(artifact)
        columns = _artifact_columns(artifact, rows_data)
        if rows_data and columns:
            artifact_table(rows_data, columns)

    if sections.get("summary", True):
        heading("1. Investigation Summary")
        two_column_table(_summary_rows(rows))
        chart = _count_chart_png("Evidence Outcome Distribution", _risk_counts(rows), THEME["blue"])
        if chart:
            document.add_picture(io.BytesIO(chart), width=Inches(6.1))

    if sections.get("evidence", True):
        heading("2. Selected Evidence Overview")
        table = document.add_table(rows=1, cols=7)
        table.style = "Table Grid"
        for index, label in enumerate(
            ["#", "Evidence", "Source", "Native verdict", "Action status", "Concern", "Engine"]
        ):
            table.rows[0].cells[index].text = label
        set_table_header(table)
        for index, row in enumerate(rows, 1):
            cells = table.add_row().cells
            values = [
                str(index),
                _evidence_family(row),
                _short_source(row)[:60],
                _native_prediction(row),
                _risk_bucket(row),
                _score_text(row),
                _engine_text(row)[:45],
            ]
            for cell_index, value in enumerate(values):
                cells[cell_index].text = value

    if sections.get("explanations", True):
        heading("3. Individual Evidence Results")
        for index, row in enumerate(rows, 1):
            family = _evidence_family(row)
            heading(f"Evidence {index} of {len(rows)} - {family}", level=2)
            two_column_table(
                [
                    ("Evidence type", family),
                    ("Source", _short_source(row)),
                    ("Native verdict", _native_prediction(row)),
                    ("Action status", _risk_bucket(row)),
                    (_text(row.get("score_label"), "Concern score"), _score_text(row)),
                    ("Engine", _engine_text(row)),
                    ("Timestamp", _date_text(row.get("scanned_at"))),
                ]
            )
            flags = ", ".join(_flags(row.get("flags")))
            if flags:
                body(f"Detected indicators: {flags}")
            explanation = _text(row.get("explanation"), "")
            if explanation:
                body(f"Explanation: {explanation}")
            preview = _text(row.get("preview"), "")
            if preview:
                body(f"Evidence preview: {preview[:350]}")
            heading("Dashboard Summary", level=2)
            two_column_table(_bundle_summary_rows(row))
            heading("Evidence Provenance", level=2)
            two_column_table(_provenance_rows(row))
            bundle = _bundle(row)
            xai = bundle.get("xai", {})
            if isinstance(xai, dict) and xai:
                heading("Explainable Analysis", level=2)
                two_column_table(
                    [
                        ("Method", _text(xai.get("method"))),
                        ("Scope", _text(xai.get("scope"))),
                        ("Explanation", _text(xai.get("explanation"))),
                    ]
                )
                factors = _xai_rows(row)
                columns = _artifact_columns({}, factors)
                if factors and columns:
                    artifact_table(factors, columns)
                for limitation in xai.get("limitations", []) if isinstance(xai.get("limitations"), list) else []:
                    body(f"XAI limitation: {_text(limitation)}")
            artifacts = bundle.get("artifacts", [])
            if isinstance(artifacts, list) and artifacts:
                heading("Saved Dashboard Artifacts", level=2)
                for artifact_index, artifact in enumerate(artifacts, 1):
                    if isinstance(artifact, dict):
                        add_artifact(artifact, artifact_index)
            else:
                body("Legacy record: no ordered dashboard artifact snapshot was saved.")

        combined = _combined_findings(rows)
        heading("Combined Investigation Findings", level=2)
        if combined:
            table = document.add_table(rows=1, cols=3)
            table.style = "Table Grid"
            for index, label in enumerate(["Finding", "Evidence involved", "Severity"]):
                table.rows[0].cells[index].text = label
            set_table_header(table)
            for item in combined:
                cells = table.add_row().cells
                cells[0].text = item["Finding"]
                cells[1].text = item["Evidence involved"]
                cells[2].text = item["Severity"]
        else:
            body("No combined indicators were available from the selected evidence.")

    if sections.get("risk", True):
        heading("4. Visual Evidence Summary")
        chart = _confidence_chart_png(rows)
        if chart:
            document.add_picture(io.BytesIO(chart), width=Inches(6.4))
        distribution_chart = _evidence_distribution_chart_png(rows)
        if distribution_chart:
            document.add_picture(io.BytesIO(distribution_chart), width=Inches(6.1))
        indicator_chart = _indicator_chart_png(rows)
        if indicator_chart:
            document.add_picture(io.BytesIO(indicator_chart), width=Inches(6.1))
        body(
            "The horizontal chart uses each source's explicitly labelled concern score where available. "
            "Model probability, audio reliability-weighted evidence, and phone rule scores remain distinct native measures."
        )

    if sections.get("recommendations", True):
        heading("5. Remediation and Prevention Plan")
        for item in _remediation_items(rows, incident_context):
            heading(f"[{item['priority']}] {item['action']}", level=2)
            body(f"Trigger: {item['trigger']}")
            body(f"Why: {item['reason']}")
        heading("Reviewer Note", level=2)
        body(report_note.strip() or DEFAULT_RECOMMENDATION)

    if sections.get("appendix", True):
        document.add_page_break()
        heading("6. Scope and Limitations")
        two_column_table(
            [
                ("System", "AI-FDS Capstone Prototype"),
                ("Email", "TF-IDF with trained email classifiers"),
                ("Transcript", "Multi-model text classification after manual input or Whisper transcription"),
                ("Audio", "MFCC voice-authenticity analysis, behavioral audio features, and transcript analysis"),
                ("Phone", "Veriphone.io carrier metadata, PenipuMY reputation evidence, and transparent rules"),
                ("Scope", "Educational scam awareness support. Not enterprise security, telecom verification, or legal evidence."),
            ]
        )
        heading("Documentation References", level=2)
        for title, url in DOCUMENTATION_REFERENCES:
            body(f"{title}: {url}")

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue(), _filename("docx")


def build_report(
    report_format: Literal["TXT", "PDF", "DOCX"],
    rows: list[dict[str, object]],
    report_note: str,
    sections: dict[str, bool],
    incident_context: dict[str, object] | None = None,
) -> tuple[bytes, str, str]:
    """Build a report and return bytes, filename, and MIME type."""

    if report_format == "TXT":
        payload, filename = build_txt(rows, report_note, sections, incident_context)
        return payload, filename, "text/plain"
    if report_format == "PDF":
        payload, filename = build_pdf(rows, report_note, sections, incident_context)
        return payload, filename, "application/pdf"
    if report_format == "DOCX":
        payload, filename = build_docx(rows, report_note, sections, incident_context)
        return payload, filename, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    raise ValueError(f"Unsupported report format: {report_format}")
