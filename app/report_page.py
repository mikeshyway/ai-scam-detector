"""AI report generator page."""

from __future__ import annotations

from io import BytesIO
from pathlib import Path

import pandas as pd
import streamlit as st

from app.ui_components import (
    render_analysis_ready,
    render_content_card_close,
    render_content_card_open,
    render_info_banner,
    render_section_header,
)
from src.time_utils import formatted_now


def _report_rows(history: list[dict[str, object]]) -> list[dict[str, object]]:
    return history[:5]


def _build_plain_report(rows: list[dict[str, object]], notes: str) -> str:
    lines = [
        "AI-based Spam and Caller Fraud Detection System",
        "Scam Analysis Report",
        f"Generated: {formatted_now()}",
        "",
        "Summary",
        f"Activities included: {len(rows)}",
        "",
        "Evidence",
    ]
    for row in rows:
        lines.extend(
            [
                f"- Time: {row.get('time', '-')}",
                f"  Type: {row.get('type', '-')}",
                f"  Prediction: {row.get('prediction', '-')}",
                f"  Confidence: {row.get('confidence', '-')}",
                f"  Evidence preview: {row.get('preview', '-')}",
            ]
        )
    lines.extend(["", "Recommendations", notes or "Verify through official channels, do not share OTP/passwords, and report suspicious activity."])
    return "\n".join(lines)


def _build_pdf(rows: list[dict[str, object]], notes: str) -> bytes | None:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
    except Exception:
        return None

    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, title="Scam Analysis Report")
    styles = getSampleStyleSheet()
    story = [
        Paragraph("AI-based Spam and Caller Fraud Detection System", styles["Title"]),
        Paragraph("Scam Analysis Report", styles["Heading2"]),
        Paragraph(f"Generated: {formatted_now()}", styles["Normal"]),
        Spacer(1, 12),
    ]
    for row in rows:
        story.extend(
            [
                Paragraph(f"{row.get('type', '-')}: {row.get('prediction', '-')}", styles["Heading3"]),
                Paragraph(f"Confidence: {row.get('confidence', '-')}", styles["Normal"]),
                Paragraph(f"Evidence preview: {row.get('preview', '-')}", styles["Normal"]),
                Spacer(1, 8),
            ]
        )
    story.extend([Paragraph("Recommendations", styles["Heading2"]), Paragraph(notes or "Verify through official channels and report suspicious activity.", styles["Normal"])])
    doc.build(story)
    return buffer.getvalue()


def _build_docx(rows: list[dict[str, object]], notes: str) -> bytes | None:
    try:
        from docx import Document
    except Exception:
        return None

    document = Document()
    document.add_heading("AI-based Spam and Caller Fraud Detection System", 0)
    document.add_heading("Scam Analysis Report", level=1)
    document.add_paragraph(f"Generated: {formatted_now()}")
    for row in rows:
        document.add_heading(f"{row.get('type', '-')}: {row.get('prediction', '-')}", level=2)
        document.add_paragraph(f"Confidence: {row.get('confidence', '-')}")
        document.add_paragraph(f"Evidence preview: {row.get('preview', '-')}")
    document.add_heading("Recommendations", level=1)
    document.add_paragraph(notes or "Verify through official channels and report suspicious activity.")
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def render_report_page(root: Path, history: list[dict[str, object]]) -> None:
    render_section_header(
        "AI analysis report generator",
        "Package recent session evidence into a portable summary for review or presentation.",
        "Evidence export",
    )
    render_info_banner(
        "Generate a lightweight evidence summary from recent session activity. "
        "Downloadable reports are used instead of SMTP to keep the capstone demonstration reliable.",
        kind="info",
        code="REPORT",
    )

    rows = _report_rows(history)
    render_content_card_open("violet")
    if rows:
        st.dataframe(pd.DataFrame(rows), hide_index=True, use_container_width=True)
    else:
        st.info("Run a detection, simulation, quiz, or phone check first. A demo report can still be generated.")

    notes = st.text_area(
        "Recommendations to include",
        value="Verify through official channels, do not share OTP/passwords, preserve evidence, and report suspicious activity to campus IT/security.",
        height=130,
    )
    render_content_card_close()

    report_rows = rows or [
        {
            "time": formatted_now(),
            "type": "Demo",
            "prediction": "Synthetic report example",
            "confidence": 0,
            "preview": "No real session evidence yet.",
        }
    ]
    plain = _build_plain_report(report_rows, notes)
    render_content_card_open("green")
    st.download_button("Download TXT report", plain, "scam_analysis_report.txt", "text/plain")

    pdf = _build_pdf(report_rows, notes)
    if pdf:
        st.download_button("Download PDF report", pdf, "scam_analysis_report.pdf", "application/pdf")
    else:
        st.warning("PDF generation requires reportlab. Install requirements.txt to enable it.")

    docx = _build_docx(report_rows, notes)
    if docx:
        st.download_button(
            "Download DOCX report",
            docx,
            "scam_analysis_report.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    else:
        st.warning("DOCX generation requires python-docx. Install requirements.txt to enable it.")
    render_content_card_close()

    if st.button("Add report generation to session history", use_container_width=True):
        history.insert(
            0,
            {
                "time": formatted_now(),
                "type": "Report",
                "prediction": "Generated",
                "confidence": 100,
                "model": "AI Report Generator",
                "preview": f"{len(report_rows)} evidence item(s)",
            },
        )
        render_analysis_ready("Report generation recorded in session history")
